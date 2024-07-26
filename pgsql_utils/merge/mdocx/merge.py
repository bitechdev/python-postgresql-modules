from ..exception import MergeToolWarning, MergeToolError
from ...util.object import getProp
from ...util.json import jsoncomplex, BJSONEnc
from ...util.string import (
    findWordArticle,
    html_replace,
    n2w,
    num2words,
    ordinalstr,
    istrreplace,
)
from ...util.number import ndec, nint, isFloat, tryInt, tryFloat, getNumber
from ...util.array import avg_any, max_any, min_any, sum_any
from ...pgproc.debug import sqlmsg, get_err_msg
from ..common.base import MergeTool
from ..taginterator import TagIterator

import xml.etree.ElementTree as ET
import lxml.etree as LET
import base64
import ast
import re
import locale


from docx import Document
from docx.shared import RGBColor
from docx.shared import Inches, Pt
import docx.text
import docx.table
import docx.enum
import docx.oxml


import io, sys, os, traceback
import json
import copy
import datetime
import math

INCH_TO_PIXEL = 0.0104166667


class MergeToolDocX(MergeTool):
    """Bitech DocX Merge Tool"""

    def __init__(s):
        super().__init__()

        s.srcdoc = None
        s.destdoc = None
        s.srcfilename = None
        s.destfilename = None
        s.visit_tables_cnt = 0
        s.tbl_basic_replace_cnt = 0
        s.tbl_map = {}
        s.tbl_ordinaldata = {}
        s.tbl_rmrow = []
        s.tbl_rmcol = []
        s.tbl_complex_rowdat = {}
        s.tbl_mode = {}
        s.tbl_search = {}

        s.white_style = None
        s.error_style = None
        s.cleanupList = []
        s.results = None

    #

    def execute(s):
        if s.mode == "peekall_docx":
            tagdata = s.peek_docx(True)
            s.results = (
                json.dumps(tagdata, ensure_ascii=False).encode("utf8").decode("utf8")
            )
        #
        elif s.mode == "peek_docx":
            tagdata = s.peek_docx()
            s.results = (
                json.dumps(tagdata, ensure_ascii=False).encode("utf8").decode("utf8")
            )
        #
        elif s.mode == "convert_docx":
            s.convert_docx()
        elif s.mode == "convertonly_docx":
            s.convertonly_docx(s.TAG_S_SP, s.TAG_E_SP)

        elif s.mode == "merge_docx":
            s.merge_docx()
        elif s.mode == "fixanchor_docx":
            s.fixanchor_to_inline(s.srcfilename)
        elif s.mode == "mocktest_docx":
            s.generate_mockup()
        elif s.mode == "mark_docx":
            s.mark_docx()
        else:
            raise MergeToolError(
                "Mode {} not supported.".format(s.mode), "DocXMergeTool.init"
            )
        #
        return s.results

    def tbl_map_byid(s, tblid):
        for nm in s.tbl_map:
            if s.tbl_map[nm].get("nr", "") == str(tblid):
                return s.tbl_map[nm]
            #
        #
        return None

    #

    def setFontProperties(s, p_src, p_dest):
        # p_dest.size = p_src.size

        try:
            p_dest.size = p_src.size
            p_dest.bold = p_src.bold
            p_dest.color.rgb = p_src.color.rgb

            p_dest.name = p_src.name
            p_dest.shadow = p_src.shadow
            p_dest.strike = p_src.strike

            # p_dest.color.type = p_src.color.type
            # p_dest.color.theme_color = p_src.color.theme_color

        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
            p_retval = 1
            p_errdetail = str(traceback.format_tb(exc_tb, 30))
            p_context = ", at line " + str(exc_tb.tb_lineno)

            sqlmsg(
                "setFontProperties failed: {} {} {} {}".format(
                    e, p_context, p_errdetail, exc_obj
                ),
                p_type="local notice",
            )
        #

    #

    def getHeadFootSections(s, p_type="head"):
        sections = []
        for sect in s.srcdoc.sections:
            if sect is None:
                continue

            if p_type.lower() == "head":
                if sect.header is not None:
                    sections.append(sect.header)
                elif sect.even_page_header is not None:
                    sections.append(sect.even_page_header)
                if sect.first_page_header is not None:
                    sections.append(sect.first_page_header)
                #
            else:
                if sect.footer is not None:
                    sections.append(sect.footer)
                elif sect.even_page_footer is not None:
                    sections.append(sect.even_page_footer)
                elif sect.first_page_footer is not None:
                    sections.append(sect.first_page_footer)
                #

            #
        #
        return sections

    #

    def actionHeadFootTags(
        s, p_type="head", p_action="list", p_fields=None, p_tables=False
    ):
        tags = []

        def checknodes(n):
            if p_tables:

                if n.tables is not None:
                    if p_action.lower() == "merge":
                        if len(n.tables) > 0:
                            s.tbl_basic_replace(n.tables)
                            s.tbl_complex_init()
                            s.tbl_post_loop(n.tables)
                        #
                        # sqlmsg("actionHeadFootTags action={} len={}".format(p_action, len(n.tables)), p_type="debug")
                    elif p_action.lower() == "list":
                        tabletags = s.visit_tables(n.tables)
                        # sqlmsg("actionHeadFootTags:tabletags: {}".format(tabletags))
                        tags.append(tabletags)
                #
            #

            #
            else:
                for p in n.paragraphs:
                    if len(p.runs) > 0:
                        for r in p.runs:
                            if p_action.lower() == "merge":
                                opdata = s.tag_value_replace(r, p, None, p_fields)
                                s.tag_picture_add(r)
                                # opdata = s.str_value_replace(r.text, None, p_fields)
                                # r.text = opdata["str"]
                            elif p_action.lower() == "convert":
                                r.text = r.text.replace(s.TAG_S_SP, s.TAG_S)
                                r.text = r.text.replace(s.TAG_E_SP, s.TAG_E)
                            elif p_action.lower() == "list":
                                tags.extend(s.peek_str(r.text, True))
                            #
                        #
                    else:
                        if p_action.lower() == "merge":
                            opdata = s.tag_value_replace(p, None, None, p_fields)
                            # opdata = s.str_value_replace(p.text, None, p_fields)
                            # p.text = opdata["str"]
                        elif p_action.lower() == "convert":
                            p.text = p.text.replace(s.TAG_S_SP, s.TAG_S)
                            p.text = p.text.replace(s.TAG_E_SP, s.TAG_E)
                        elif p_action.lower() == "list":
                            tags.extend(s.peek_str(p.text, True))
                        #
                    #
                #
            #

        #

        try:
            sections = s.getHeadFootSections(p_type)
            # sqlmsg("actionHeadFootTags get={} sections={} -> {}".format(p_action, len(sections), sections), p_type="debug")
            for sect in sections:
                checknodes(sect)
            #
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
            m_errmsg = "{} Trace: {}, Line: {}".format(
                str(e), str(traceback.format_tb(exc_tb, 30)), str(exc_tb.tb_lineno)
            )

            sqlmsg(
                "actionHeadFootTags Exception {}".format(m_errmsg), p_type="local error"
            )
        #
        return tags

    #

    def load_docx(s):
        if s.srcfilename is None or len(s.srcfilename) < 3:
            stream = io.BytesIO(s.template_blob)
            s.srcdoc = Document(stream)
        else:
            s.srcdoc = Document(s.srcfilename)
        #

    #

    def save_docx(s):

        if s.destfilename is None or len(s.destfilename) < 3:
            stream = io.BytesIO()
            s.srcdoc.save(stream)
            s.results = stream.getvalue()
        else:
            s.srcdoc.save(s.destfilename)
            if sys.platform.lower() == "linux":
                os.chmod(s.destfilename, 0o777)
            #
        #
        return s.results

    #

    def add_table_tree(s, treedata, p_parentcnt=0):
        tags = {}

        for dat in treedata:
            if dat["parsed"]:
                s.visit_tables_cnt = s.visit_tables_cnt + 1
                tablekey = "{},{}".format(s.visit_tables_cnt, p_parentcnt)
                tags[tablekey] = dat["parsed"]
            #
            if dat.get("children") and len(dat["children"]) > 0:
                s.add_table_tree(dat["children"], s.visit_tables_cnt)
            #
        #

        return tags

    #

    def peek_alias_str(s, p_str, p_tablename, p_mode, p_obj, p_datIndex=None):
        """Scans through a string for merge aliases and build a table with it."""
        alias_tags = []
        if p_str is None or p_str == "":
            return p_str
        objtype = ""
        if isinstance(p_obj, docx.text.run.Run):
            objtype = "run"
        #
        if len(s.aliasList) > 0:
            for r in s.aliasList:
                if r["obj"] == p_obj and r["type"] == "declare":
                    return alias_tags
                #
            #
        #

        def search_str(p_str, p_tag_s, p_tag_e, p_start=0):
            # global tags
            i_s = p_str.find(p_tag_s, p_start)
            i_e = p_str.find(p_tag_e, i_s)
            if i_s >= 0 and i_e > i_s:
                tagbody = p_str[i_s + len(p_tag_s) : i_e]
                tagname = ""
                ops = []
                if len(tagbody) > 1:
                    opsplitPos = tagbody.find(":")
                    if opsplitPos >= 0 and p_tag_s != s.TAG_S_ALIASDECL:
                        # ops tags only work in use situations
                        # tagbody = tagbody[:opsplitPos]
                        tagname = str(tagbody[:opsplitPos]).lstrip().rstrip()
                        ops = str(tagbody[opsplitPos + 1 :].lstrip().rstrip()).split(
                            ":"
                        )
                    #
                    fulltag = p_tag_s + tagbody + p_tag_e
                    if tagname == "":
                        tagname = tagbody
                    #
                    record = {
                        "type": "use",
                        "tag": fulltag,
                        "body": tagbody,
                        "name": tagname,
                        "table": p_tablename,
                        "declid": None,
                        "obj": p_obj,
                        "objtype": objtype,
                        "ops": ops,
                        "datindex": 0,
                    }

                    if p_tag_s == s.TAG_S_ALIASDECL:
                        parts = str(tagbody).split(s.TAG_D_ALIASDECL, 1)
                        record["type"] = "declare"
                        record["name"] = str(parts[0]).lstrip().rstrip()
                        record["value"] = parts[1]
                        record["datindex"] = p_datIndex

                    #
                    alias_tags.append(record)
                    # recursive to next position.
                    search_str(p_str, p_tag_s, p_tag_e, i_e + len(p_tag_e))
                    return p_str
                #
            #
            return ""

        #
        res = None
        res1 = None
        if p_mode == "peekuse":
            res1 = search_str(p_str, s.TAG_S_ALIASUSE, s.TAG_E_ALIASUSE, 0)
            if res1 == "":
                return
        elif p_mode == "peekdeclare":
            res = search_str(p_str, s.TAG_S_ALIASDECL, s.TAG_E_ALIASDECL, 0)
            if res == "":
                return
        else:
            res = search_str(p_str, s.TAG_S_ALIASDECL, s.TAG_E_ALIASDECL, 0)
            res1 = search_str(p_str, s.TAG_S_ALIASUSE, s.TAG_E_ALIASUSE, 0)
            if res == "" and res1 == "":
                return
        #
        if p_mode == "peek":
            if s.aliasListPeek is None:
                s.aliasListPeek = []
            #
            tags = []
            for t in alias_tags:
                tags.append(t["tag"])
            #
            s.aliasListPeek.extend(tags)
            # sqlmsg(json.dumps(s.aliasListPeek), "aliasListPeek")
        else:
            s.aliasList.extend(alias_tags)
            # sqlmsg(json.dumps(s.aliasListPeek), "aliasList")
        #
        # sqlmsg("peek_alias_str {}: alias_tags:{} {} p_str:{}".format(p_mode,alias_tags,s.aliasList,p_str), p_title="local error")

        # s.buidAliasMap()

        return alias_tags

    #

    def peek_docx(s, p_inc_old=False):
        tags = {}
        try:
            if s.srcdoc is None:
                s.load_docx()
            #

            doctags = []
            tags["doctags"] = []
            tags["pictures"] = []
            tags["headertags"] = []
            tags["footertags"] = []
            tags["shapes"] = []
            tags["sections"] = []
            tags["raw"] = []
            tags["aliases"] = []
            tags["tabletags"] = {}

            txtTags = TagIterator()
            txtTags.start()

            def fnProcess(txt):
                return s.peek_str(txt)

            #
            txtTags.fnProcess = fnProcess

            for p in s.srcdoc.paragraphs:
                # plpy.notice("Reading paragraph: {}",format(p.text)) ##debug,remove
                tags["doctags"].extend(s.peek_str(p.text))
                s.peek_alias_str(p.text, "", "peek", p)

                txtTags.process_text(p.text)

                # tags["tabletags"]
                # for r in p.runs:
                ##depreciated, using mergetag pic replace.
                # for sp in r.part.inline_shapes:
                #  if sp.type ==  docx.enum.shape.WD_INLINE_SHAPE.PICTURE:
                #    #plpy.notice("Pic: {}".format(sp))
                #    tags["pictures"].append({ "filename": os.path.basename(sp._inline.docPr.get('descr',''))
                #                             ,"name": sp._inline.docPr.get('name','')
                #                             })
                ##
                #

            #

            for sec in s.srcdoc.sections:
                tags["sections"].append(str(sec.start_type))
            #

            for x in s.srcdoc.inline_shapes:
                tags["shapes"].append(str(x.type))
            #

            tags["tabletags"].update(s.visit_tables(s.srcdoc.tables))
            tags["tabletags"].update(txtTags.as_merge_tags(s.visit_tables_cnt + 1))

            tags["headertags"].extend(s.actionHeadFootTags("head", "list", None, False))
            tags["footertags"].extend(s.actionHeadFootTags("foot", "list", None, False))

            moretableTags = s.actionHeadFootTags("head", "list", None, True)
            moretableTags.extend(s.actionHeadFootTags("foot", "list", None, True))
            for a in moretableTags:
                tags["tabletags"].update(a)
            #

            tags["raw"].extend(s.docx_textboxes_lp("peek"))
            tags["aliases"].extend(s.aliasListPeek)

            # plpy.notice("Total Shapes: {}".format(len(s.srcdoc.inline_shapes)))
            # depreciated, using mergetag pic replace.
            # for sp in s.srcdoc.inline_shapes:
            #  if s.type ==  docx.enum.shape.WD_INLINE_SHAPE.PICTURE:
            #    try:
            #      tags["pictures"].append({ "filename": os.path.basename(os.path.normpath(sp._inline.docPr.get('descr','')))
            #                               ,"name": sp._inline.docPr.get('name','')
            #                               })
            #    except Exception as e:
            #      plpy.warning("Could not load shape. Details: {}".format(e))
            #    #
            ##
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            raise MergeToolError("Failed reading docx: " + str(e), "peek_docx", exc_tb)
        #
        return tags

    #

    def docx_textboxes_lp(s, p_type):
        tags = []
        body = s.srcdoc._body
        text_box_p_elements = body._body.xpath(".//w:txbxContent//w:t")
        # plpy.notice("Found : {}".format(text_box_p_elements))
        for xm in text_box_p_elements:
            if p_type == "merge":
                opdata = s.str_value_replace(str(xm.text))
                # sqlmsg("Replaced: {}".format(xm.text))
                xm.text = opdata["str"]

                # if opdata["error"] is None:
                #  sqlmsg("With: {}".format(xm.text))
                # else:
                #  sqlmsg("Error: {}".format(opdata["error"]))
                ##
            else:
                tags.extend(s.peek_str(xm.text))
            #
        #

        return tags

    #

    def visit_table_cb(s, p_tables, p_parent=0, p_fn_callback=None):
        table_tags = {}
        if p_parent == 0:
            s.visit_tables_cnt = 0

        for t in p_tables:
            s.visit_tables_cnt += 1
            tablekey = "{},{}".format(s.visit_tables_cnt, p_parent)
            table_tags[tablekey] = []
            # plpy.notice("X{}".format("{},{}".format(s.visit_tables_cnt,p_parent)))
            lastid = s.visit_tables_cnt
            rcnt = 0
            for r in t.rows:
                rcnt += 1
                for c in r.cells:
                    for p in c.paragraphs:
                        if len(p.runs) > 0:
                            for pr in p.runs:
                                p_fn_callback(tablekey, t, "run", pr, rcnt, r, c)
                            #
                        else:
                            p_fn_callback(tablekey, t, "paragraph", p, rcnt, r, c)
                        #
                    #
                    if len(c.tables) > 0:
                        tags = s.visit_table_cb(c.tables, lastid, p_fn_callback)
                        for t in tags:
                            table_tags[t] = tags[t]
                        #
                    #
                #
            #

        #
        return table_tags

    #

    def visit_tables(s, p_tables, p_parent=0):

        table_tags = {}
        if p_parent == 0:
            s.visit_tables_cnt = 0

        txtTags = TagIterator()
        txtTags.start()

        def fnProcess(txt):
            return s.peek_str(txt)

        #
        txtTags.fnProcess = fnProcess

        for t in p_tables:
            if not s.unique_table_check(t._tbl):
                continue
            #
            s.visit_tables_cnt += 1
            tablekey = "{},{}".format(s.visit_tables_cnt, p_parent)
            # table_tags[tablekey] = []
            # plpy.notice("X{}".format("{},{}".format(s.visit_tables_cnt,p_parent)))
            lastid = s.visit_tables_cnt
            rcnt = 0
            for r in t.rows:
                rcnt += 1
                for c in r.cells:
                    for p in c.paragraphs:
                        # lst = s.peek_str(p.text)
                        txtTags.process_text(p.text)
                        s.peek_alias_str(p.text, tablekey, "peek", p, rcnt)
                        # if len(lst) > 0:
                        #   if table_tags.get(tablekey) is None:
                        #     table_tags[tablekey] = []
                        #   #
                        #   table_tags[tablekey].extend(lst)
                        # #
                    #
                    if len(c.tables) > 0:
                        tags = s.visit_tables(c.tables, lastid)
                        for t in tags:
                            table_tags[t] = tags[t]
                        #
                    #
                #
            #
        #

        # plpy.notice("visit_tables_cnt: {}\n".format(s.visit_tables_cnt))

        table_tags.update(
            txtTags.as_merge_tags(s.visit_tables_cnt + 1, s.visit_tables_cnt, True)
        )
        # plpy.notice("\ntable_tags: {}\n".format(table_tags))

        return table_tags

    #

    # def istrreplace(p_str, p_value, p_newval, p_cnt = 0):
    #  if type(p_str) is not str: return p_str
    #  if p_value.find(p_newval) > 0: return p_str
    #  istart = p_str.lower().find(p_value.lower())
    #  if istart < 0: return p_str
    #  vallen = len(p_value)
    #  if istart >= 0:
    #    p_value = p_str[:istart] + p_newval + p_str[:istart+vallen]
    #  #
    #  if p_cnt > 100: return p_str #we only support 100 interations, in case you try and replace the same str.
    #  p_str = istrreplace(p_str, p_value, p_newval, p_cnt+1)
    #  return p_str
    #
    def run_add_run(s, p_run, p_parent, p_text):
        p_xml = p_run._parent
        r_xml = p_xml._p.add_r()
        # sqlmsg("run_add_run {} {}".format(p_run, p_parent))
        # r_xml = docx.oxml.text.run.CT_R()
        newrun = docx.text.run.Run(r_xml, p_parent)
        newrun.text = p_text
        if p_run.style is not None:
            newrun.style = p_run.style
        #
        # if p_parent.style is not None:
        #  newrun.style = p_parent.style
        #
        # p_run._r.addnext(newrun._r)
        return newrun

    #

    def docxcitation(s, p_obj, p_parent, p_val):
        # return False
        try:
            valjson = json.loads(p_val)
            if valjson and valjson["mergetype"]:
                # sqlmsg("Found mergetype {}".format(valjson))
                line = valjson["line"]
                for k, v in valjson.items():
                    valjson["line"] = valjson["line"].replace(k, v)
                    # sqlmsg("valjson: {}".format(valjson["line"]))
                #

                bolds = valjson["line"].split("<")
                line = valjson["line"]
                for i in range(len(bolds) + 1):
                    nodestart = line.find("<")
                    if nodestart < 0:
                        break
                    nodeend = line.find(">")
                    node = line[nodestart + 1 : nodestart + 2]
                    if node in ("a", "b"):
                        endtag = "</{}>".format(node)
                        nodeendtag = line.find(endtag)
                        partstart = line[:nodestart]
                        partmid = line[nodeend + 1 : nodeendtag]
                        line = line[nodeendtag + len(endtag) :]

                        # sqlmsg("s:{} e:{} node {}; endtag: {};".format(nodestart,nodeend, node, nodeendtag))
                        # sqlmsg("partstart:{} partmid {}".format(partstart, partmid))

                        # startobj = p_obj.add_run(partstart)

                        startobj = s.run_add_run(p_obj, p_parent, partstart)
                        startobj.style = p_obj.style
                        s.setFontProperties(p_obj.font, startobj.font)

                        # midobj = p_obj.add_run(partmid)
                        midobj = s.run_add_run(p_obj, p_parent, partmid)
                        midobj.style = p_obj.style
                        s.setFontProperties(p_obj.font, midobj.font)

                        # midobj.font = font
                        # (str(type(p_obj)).find("Run") >= 0
                        if node == "b":
                            midobj.bold = True
                        #
                        if node == "a":
                            midobj.underline = True
                            font = midobj.font
                            font.color.rgb = RGBColor(46, 100, 152)
                        #

                    else:
                        # sqlmsg("Continue s:{} e:{} node {}; line: {}".format(nodestart,nodeend, node, line))
                        continue
                    #

                    # sqlmsg("Start {} End: {} line:{}".format(start,end, line))
                #

            #
        # namepart = p_obj.add_run(valjson["line"])
        except Exception as e:
            sqlmsg("docxcitation error: {}".format(e), p_type="local error")
            return False
        #
        return True

    #

    def tag_value_replace(
        s, p_obj, p_parent, p_idx=None, p_fields=None, p_type="", p_tblid="", p_colidx=0
    ):
        """Scans through docx objects and replace from the given values."""
        opdata = {
            "error": None,
            "rmblnk": 0,
            "rmdr": 0,
            "mergedtags": [],
            "value": None,
            "rmblnkcol": 0,
            "limitrow": 0,
            "mergedvalues": [],
        }
        isrun = False
        try:

            if p_idx is None:
                p_idx = 0

            if p_obj is None:
                opdata["error"] = "No page object. "
                return opdata
            #

            if not (
                str(type(p_obj)).find("Run") >= 0
                or str(type(p_obj)).find("Paragraph") >= 0
            ):
                raise Exception("Only run and Paragraph objects supported.")
            #
            isrun = str(type(p_obj)).find("Run") >= 0

            if p_obj.text is None:
                opdata["error"] = "No tags found."
                return opdata
            #

            s.peek_alias_str(p_obj.text, p_tblid, "peekuse", p_obj, p_idx)

            if (
                not p_obj.text.find(s.TAG_S) >= 0
            ):  # abort this line/object, we have not tags.
                opdata["error"] = "No tags found."
                return opdata
            #
            fields = p_fields
            if fields is None:
                fields = s.inputdata.get("fields")
            #

            if p_obj.text.find(s.TAG_S + s.STR_TAG_RMBD + s.TAG_E) >= 0:
                opdata["rmblnk"] = 1
            #

            if p_obj.text.find(s.TAG_S + s.STR_TAG_RMBC + s.TAG_E) >= 0:
                opdata["rmblnkcol"] = 1
            #

            if p_obj.text.find(s.TAG_S + s.STR_TAG_RMDR + s.TAG_E) >= 0:
                opdata["rmdr"] = 1
            #

            if p_obj.text.find(s.TAG_S + s.STR_TAG_LIMITROW) >= 0:
                istart = p_obj.text.find(s.TAG_S + s.STR_TAG_LIMITROW)
                iend = p_obj.text.find(s.TAG_E, istart)
                parts = (p_obj.text[istart + len(s.TAG_S) : iend]).split(":")
                if len(parts) > 1:
                    opdata["limitrow"] = tryInt(parts[1])
                else:
                    opdata["limitrow"] = 1
                #
                # sqlmsg("Table limitrow: {} -> {} opdata:{} ".format(p_obj.text,parts, opdata ))
                p_obj.text = p_obj.text[:istart] + p_obj.text[iend + len(s.TAG_E) :]
            #

            startSearch = p_obj.text.find(s.TAG_S + s.STR_TAG_TBLSEARCH_RM)
            if startSearch >= 0:
                firstpartLen = len(s.TAG_S + s.STR_TAG_TBLSEARCH_RM)
                startEnd = p_obj.text.find(s.TAG_E, startSearch + firstpartLen)
                if startEnd >= 0:
                    foundtext = str(p_obj.text)[startSearch + firstpartLen : startEnd]
                    opvalues = foundtext.split(":")

                    if len(opvalues) > 1:
                        s.tbl_search[p_tblid] = {
                            "searchvalue": ":".join(opvalues[1:]),
                            "colindex": p_colidx,
                            "rowindex": p_idx,
                            "objptr": p_obj,
                            "action": "remove",
                            "found": False,
                        }

                    #
                #
            #

            if fields is None or type(fields) is not dict:
                # tags = s.peek_str(p_obj.text,False)
                # plpy.warning("No fields to replace in tag_value_replace. {} {}".format(fields,tags))
                opdata["error"] = "No fields given."
                return opdata
            #

            lasttag = None
            for tag in fields:
                found = False
                lasttag = tag
                val = None
                tagtype = None
                tagvalue = None

                try:
                    if type(fields.get(tag)) is not dict:
                        raise MergeToolError(
                            "Expecting tag field to have a type and value property.",
                            "tag_value_replace",
                        )
                    #
                    if p_obj.text is None:
                        break
                    if (
                        not p_obj.text.find(s.TAG_S) >= 0
                    ):  # abort this line/object, we have not tags.
                        # plpy.notice('No more tags, last check for {}'.format(tag));
                        break  # this is a funky break to exit 2nd attempt. Works with check above, don't remove.
                    #

                    tagtype = int(fields.get(tag).get("type", 0))
                    tagvalue = fields.get(tag).get("value", "")

                    if (
                        type(tagvalue) is list
                    ):  # we can look at the type also, but sigle values wont have a list
                        try:
                            val = tagvalue[p_idx]
                            found = True
                        except Exception as e:
                            s.add_fault("No data for field {}".format(tagtype), str(e))
                            plpy.warning(
                                "Field replace Exception: {} Detail: i:{} tag:{},datalist:{}".format(
                                    e, p_idx, tagtype, tagvalue
                                )
                            )
                            pass
                        #
                    else:
                        val = tagvalue
                        if tagvalue != "":
                            found = True
                        #
                    #
                    # opdata["mergedvalues"].append({"col": p_colidx,"tag": tag, "value":val, "found": found})

                    ### E-Sign tags
                    if (
                        p_obj.text.lower().find(s.STR_TAG_ESIGN) >= 1
                        and tag.lower().find(s.STR_TAG_ESIGN) >= 0
                    ):
                        formattype = "esign"
                        # sqlmsg("Found E-Sign Tag: {} val:{} T:{}".format(tag,val, p_obj.text))
                        # if isinstance(p_obj, docx.text.paragraph.Paragraph):
                        #  p_obj.text = p_obj.text.replace(tag, "")
                        #  run = p_obj.add_run(val)
                        #  s.post_format.append({"tag": tag, "obj": run, "type": "esign"})
                        # else:
                        if val is not None and val.find("!skip") >= 0:
                            plpy.notice("Esign Tag is skipped, {}".format(tag))
                            # these are used by docminer and must stay.
                        elif val is not None and tagtype == s.TTYP_PIC:
                            plpy.notice(
                                "Esign Tag is a picture,leave it be for picture replace, {}".format(
                                    tag
                                )
                            )
                            s.cleanupList.append({"obj": p_obj, "tag": tag})
                        else:
                            p_obj.text = p_obj.text.replace(tag, val)
                            if val is not None and val.find("!") < 0:
                                formattype = "remove"
                            #
                        #
                        # run = p_obj.add_run(val,white_style)
                        s.post_format.append(
                            {"tag": tag, "obj": p_obj, "type": formattype}
                        )

                        continue
                    ##
                    if tagtype == s.TTYP_PIC:
                        plpy.notice("Picture Tag, {}  tagtype: {}".format(tag, tagtype))
                        # s.cleanupList.append({"obj": p_obj, "tag":tag })
                        continue
                    #
                    if tagtype == s.TTYP_TBLROOT:
                        continue  # we do not handle inner values here.
                    if tagtype == s.TTYP_CONDITIONAL:
                        if p_obj.text.find(tag) >= 0:
                            if tryInt(tagvalue) > 0:
                                if tag.find("rmr!") >= 0:
                                    val = s.TAG_COND_RMR
                                    # sqlmsg("Will Remove Row: {} , {}, p_str={}".format(tag, tagvalue, p_obj.text))
                                elif tag.find("rmc!") >= 0:
                                    val = s.TAG_COND_RMC
                                elif tag.find("cc!") >= 0:
                                    val = s.TAG_COND_CLR
                                elif tag.find("!rmp!") >= 0:
                                    val = s.TAG_COND_RMP
                                else:
                                    continue
                                #
                                # plpy.notice("Conditional Type({}) Tag Value: {} = {}  - opdata={}".format(p_type,tag,tagvalue,opdata))
                            else:
                                continue
                            #
                        #
                    #
                    if (
                        val
                        and tagtype in (s.TTYP_FIELD, s.TTYP_TBLFIELD)
                        and val.find(s.TAG_S_INPLACE) >= 0
                        and p_obj.text.find(tag) >= 0
                    ):
                        filename = val.replace(s.TAG_S_INPLACE, "").replace(
                            s.TAG_E_INPLACE, ""
                        )
                        # p_obj.text = p_obj.text.replace(tag, "")
                        # sqlmsg("InPlace {} for {}".format(filename, p_obj.text))
                        s.placeDoc(p_obj, p_parent, tag, filename)
                        continue
                    #

                    if (
                        val
                        and tagtype in (s.TTYP_FIELD, s.TTYP_TBLFIELD)
                        and val.find("mergetype") >= 0
                        and val.find("line") >= 0
                        and p_obj.text.find(tag) >= 0
                    ):

                        if (
                            s.cleanMissing
                            and not tag.lower().find(s.STR_TAG_ESIGN) >= 0
                        ):
                            p_obj.text = p_obj.text.replace(tag, "")
                        #
                        s.docxcitation(p_obj, p_parent, val)

                        # if isrun:
                        # else:
                        #  s.docxcitation(p_obj, val)
                        #

                        break
                    #

                    if tagtype == s.TTYP_DOCREPLACE:
                        if s.cleanMissing:
                            p_obj.text = p_obj.text.replace(tag, "")
                        #
                        s.placeDoc(p_obj, p_parent, tag, val)
                        continue
                    #
                    if val is None:
                        val = ""
                        if opdata["rmblnk"] == 1:
                            opdata["rmblnk"] = 2
                        if opdata["rmblnkcol"] == 1:
                            opdata["rmblnkcol"] = 2
                    #
                    if type(val) is not str:
                        plpy.warning(
                            "Expected string type value. {}".format(type(val), val)
                        )
                        opdata["error"] = "Value not of type string"
                        break
                        # return opdata
                    #

                    if val == "" or val == "0":
                        if opdata["rmblnk"] == 1:
                            opdata["rmblnk"] = 2
                        if opdata["rmblnkcol"] == 1:
                            opdata["rmblnkcol"] = 2

                        found = False
                    #

                    if found:
                        if p_obj.text.find(tag) > -1:
                            if opdata["rmdr"] == 1:
                                opdata["rmdr"] = 2
                            #

                            opdata["mergedtags"].append(tag)
                            if opdata["value"] is None:
                                opdata["value"] = val
                            #
                        #
                    #
                    if val == "":
                        txt = p_obj.text.replace(tag, "")
                        if len(txt) <= 1:
                            p_obj.text = "."
                            p_obj.style = s.white_style
                            font = p_obj.font
                            font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            continue
                        #
                    #
                    p_obj.text = p_obj.text.replace(tag, val)
                    # plpy.notice("Tag replace: {}={}  obj:{}".format(tag, val, p_obj.text))
                except Exception as e:
                    sqlmsg(
                        "Failed to replace tag value. tag: {} tagvalue: {} Error: {}".format(
                            tag, tagvalue, e
                        ),
                        p_type="local notice",
                    )
                #
                #
            #
        except Exception as e:
            sqlmsg(
                "Failed to replace tag value. Fields: {} Idx: {} Error:{}".format(
                    p_fields, p_idx, e
                ),
                p_type="local notice",
            )
        #
        return opdata

    #

    def tag_replace_picture(s, p_obj, p_parent, p_key, p_value):
        """Replace key with a picture. Expecting base64 value"""
        filename = ""

        if str(type(p_obj)).lower().find("inlineshape") >= 0:
            filename = ""
            try:
                filename = p_obj._inline.docPr.get("descr", "")
                filename = os.path.basename(os.path.normpath(filename))
            except Exception as e:
                plpy.notice("Failed to get inline data: {}, {}".format(e, p_obj))
                return
            #
            if filename == "":
                return
            #
        #
        # plpy.notice("Picture replace -> {} {} ObjType: {} , {}".format(p_key, filename, str(type(p_obj)), p_parent ))

        try:
            imgstream = io.BytesIO(base64.b64decode(p_value))

            if str(type(p_obj)).find("Run") >= 0:
                p_obj.text = p_obj.text.replace(p_key, "")
                # p_obj.add_picture(imgstream)
            elif str(type(p_obj)).find("Paragraph") >= 0:
                p_obj.text = p_obj.text.replace(p_key, "")
                run = p_obj.add_run()
                # run.add_picture(imgstream)
            elif str(type(p_obj)).lower().find("inlineshape") >= 0:
                ifname = p_obj._inline.docPr.get("descr", "")
                iname = p_obj._inline.docPr.get("name", "")
                h = p_obj.height
                w = p_obj.width

                if ifname.lower().find(p_key.lower()) < 0:
                    plpy.notice(
                        ">>No match for imgfilename: {}, searchname: {}".format(
                            ifname, p_key
                        )
                    )
                    return
                #
                plpy.notice(
                    ">> Image W,H: {}, {} imgfilename: {}, searchname: {}".format(
                        w, h, ifname, p_key
                    )
                )
                # docx.text.run.Run(p_obj._inline.getparent().getparent(), p_obj._inline.getparent().getparent())
                if p_parent is not None and str(type(p_parent)).find("Run") >= 0:
                    p_parent.add_picture(imgstream, w, h)

                #
                p_obj = None
                # p_obj._inline.getparent().remove(p_obj._inline)
                plpy.notice(">> Success {}".format(p_key))
            #
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            raise MergeToolError(
                "Failed replacing picture into document. Technical Details:\r\n{}".format(
                    e
                ),
                "tag_replace_picture",
                exc_tb,
            )
        #

    #

    def loop_shapes(s, p_inline_shapes, p_parent):
        cnt = len(p_inline_shapes)
        # plpy.notice("Total shapes: {}".format(cnt))
        if cnt == 0:
            return

        fields = s.inputdata.get("fields")

        for tag in fields:
            if type(fields.get(tag)) is not dict:
                raise MergeToolError(
                    "Expecting tag field to have a type and value property.",
                    "tag_value_replace",
                )
            #
            tagtype = int(fields.get(tag).get("type", 0))
            if tagtype == 5:
                tagvalue = fields.get(tag).get("value", "")

                for sp in p_inline_shapes:
                    imagefilename = sp._inline.docPr.get("descr", "")
                    imagefilename = os.path.basename(imagefilename)
                    if imagefilename.find(tag) >= 0:
                        s.tag_replace_picture(sp, p_parent, tag, tagvalue)
                    #
                #
            #
        #

    #

    def replace_shapes(s, p_inline_shapes):
        cnt = len(p_inline_shapes)
        # plpy.notice("Total shapes: {}".format(cnt))
        if cnt == 0:
            return

        for sp in p_inline_shapes:
            imagefilename = sp._inline.docPr.get("descr", "")
            base64str = None
            for tag in s.inputdata.get("fields"):
                if imagefilename.lower().find(tag.lower()) < 0:
                    continue
                tagdata = s.inputdata.get("fields").get(tag)
                if type(tagdata) is not dict:
                    continue
                if int(tagdata.get("type", "0")) != 5:
                    continue
                base64str = tagdata.get("value", "")
                break
            #

            if base64str is None or len(base64str) < 4 or imagefilename == "":
                plpy.notice("Nothing found {}".format(imagefilename))
                break
            #
            h = sp.height
            w = sp.width
            imgstream = io.BytesIO(base64.b64decode(base64str))
            # sp.get_or_add_image()
            _drawing = sp._inline.getparent()
            _run = _drawing.getparent()
            _run.remove(_drawing)
            for p in s.srcdoc.paragraphs:
                for r in p.runs:
                    if r._r == _run:
                        plpy.notice("Found run: {}, placing image".format(r))
                        r.add_picture(imgstream, w, h)
                    #
                #
            #

            # plpy.notice("!!Parent: {}".format(_run))
        #

    #

    def tag_picture_add(s, p_objrun, p_fields=None):
        """Adds a picture from fields into run where the given tag is found ."""
        base64str = None

        if p_fields is None:
            p_fields = s.inputdata.get("fields")
        #

        if p_fields is None:
            return
        if p_objrun is None:
            return

        def findtagdetails(p_sts, p_type="w"):
            val = ""
            i_ste = p_objrun.text.find("^]", p_sts) + 2
            i_ws = p_objrun.text.find(p_type + "=", p_sts, i_ste)
            i_we = p_objrun.text.find(",", i_ws + 2, i_ste)
            if i_we < 0:
                i_we = p_objrun.text.find("^", i_ws + 2, i_ste + 1)
            #
            if i_we >= 0 and i_ws >= 0:
                val = p_objrun.text[i_ws + 2 : i_we]
            #

            # sqlmsg("add findtagdetails for tag:{} , p_sts:{}, i_ste:{},i_ws:{},i_we:{}  ,val:{},run:{}".format(tag,p_sts,i_ste,i_ws,i_we,val,p_objrun.text), p_title="pl_mailmerge", p_type="local notice")
            if val.isalnum():
                return float(val)
            else:
                return 0
            #

        #

        for tag in p_fields:
            try:
                if p_objrun.text.find(tag) < 0:
                    continue
                tagdata = s.inputdata.get("fields").get(tag)
                if type(tagdata) is not dict:
                    continue
                validpicture = True
                if tryInt(tagdata.get("type", "0")) != s.TTYP_PIC:
                    validpicture = False
                #
                base64str = tagdata.get("value", "")
                if base64str is None or len(base64str) < 4:
                    validpicture = False
                #
                h = tryFloat(tagdata.get("h", "0.5"), 0.5)
                w = tryFloat(tagdata.get("w", "2"), 2)
                user_w = 0
                user_h = 0

                # todo read these.

                i_ts = p_objrun.text.find(tag)
                i_sts = p_objrun.text.find("[^", i_ts)
                # sqlmsg("test findtagdetails for tag:{} , i_ts:{} , i_sts:{} ,val:{}".format(tag,i_ts,i_sts,p_objrun.text), p_title="pl_mailmerge", p_type="local notice")
                if i_sts >= 0 and i_sts >= i_ts:
                    user_w = findtagdetails(i_sts, "w")
                    user_h = findtagdetails(i_sts, "h")
                    i_ste = p_objrun.text.find("^]", i_ts)
                    p_objrun.text = p_objrun.text[:i_sts] + p_objrun.text[i_ste + 2 :]
                #

                objw = None
                objh = None

                if w != 0:
                    objw = Inches(w)
                #
                if h != 0:
                    objh = Inches(h)
                #
                if user_w > 0 or user_h > 0:
                    objw = None
                    objh = None
                    if user_w > 0:
                        objw = Inches(user_w * INCH_TO_PIXEL)
                    #
                    if user_h > 0:
                        objh = Inches(user_h * INCH_TO_PIXEL)
                    #
                #
                if validpicture:
                    imgstream = io.BytesIO(base64.b64decode(base64str))
                    p_objrun.text = p_objrun.text.replace(tag, "")
                    p_objrun.add_picture(imgstream, objw, objh)
                #
                # sqlmsg("add picture for {} , {} ,W: {}, H:{}".format(tag,tagdata,w,h), p_title="pl_mailmerge", p_type="local notice")
            except Exception as e:
                sqlmsg(
                    "Failed to add picture for {} - {} ,Error: {}".format(
                        tag, tagdata, e
                    ),
                    p_title="pl_mailmerge",
                    p_type="local notice",
                )
            #
        #

    #

    def markerr_obj(s, p_obj, p_objpar, p_values=None, p_tblsource=""):
        """Scans through a string for incomplete tags and mark them."""
        global applyError, errorstr
        applyError = False
        errorstr = ""

        # i_s = p_obj.text.lower().find(s.STR_ERR)
        # i_e = i_s+len(s.STR_ERR)
        # if i_s > 0 and len(p_obj.text) >= i_e:
        # p_obj.text = ''
        ##  sqlmsg("Found exiting err tag. Removing {}".format(p_obj.text))
        #  return None
        #

        def search_str(p_str, p_tag_s, p_tag_e, p_start=0):
            ###Inner Routine to search and add faults
            global applyError, errorstr

            slen = len(p_str)
            i_s = p_str.lower().find(p_tag_s, p_start)
            i_e = p_str.lower().find(p_tag_e, i_s)

            if i_s >= 0 and p_str[i_s + 1 : i_s + 1 + len(s.TAG_ERR)] == s.TAG_ERR:
                fulltag = p_str[i_s : i_e + len(p_tag_e)]
                s.add_fault("not found", fulltag)

                if s.DEBUG:
                    plpy.notice(
                        "Already error checked. {}".format(
                            p_str[i_s : i_e + len(p_tag_e)]
                        )
                    )
                #
                return p_str
            elif i_s >= 0 and i_e > i_s:
                fulltag = p_str[i_s : i_e + len(p_tag_e)]
                found = False
                if i_e - i_s > 100:
                    applyError = True
                    s.add_fault("Tag to long. More than 100 characters.", fulltag)
                #
                if p_values is not None:
                    for v in p_values:
                        errorstr = v.get("error", "Not found")
                        tblsrc = v.get("source", "")

                        if (
                            fulltag != ""
                            and v.get("tag", "").find(fulltag) >= 0
                            and (
                                tblsrc == p_tblsource
                                or p_tblsource == ""
                                or tblsrc == ""
                            )
                        ):
                            found = True
                            # plpy.notice("Found - {} -> {} Src: {}:{}".format(fulltag, v, p_tblsource, tblsrc))
                        #
                    #
                    if found:
                        # p_str = p_str[:i_s+len(p_tag_s)] + s.TAG_ERR + p_str[i_s+len(p_tag_s):]
                        p_str = p_str  # we are not replacing it anymore. Just add error below.
                        applyError = True
                        # s.add_fault("not found", fulltag)
                    #
                # plpy.notice("Cheking tag: {} found: {} Started@{}".format(fulltag,found,p_start))
                p_str = search_str(
                    p_str, p_tag_s, p_tag_e, i_e + len(p_tag_e) + (len(p_str) - slen)
                )
            elif i_s >= 0:
                if s.DEBUG:
                    plpy.notice(
                        "Broken tag: {} Started@{}".format(
                            p_str[i_s + len(p_tag_s) : i_s + len(p_tag_s) + 30], p_start
                        )
                    )
                #
                p_str = (
                    p_str[: i_s + len(p_tag_s)]
                    + s.TAG_ERR
                    + p_str[i_s + len(p_tag_s) :]
                )
                p_str = search_str(
                    p_str,
                    p_tag_s,
                    p_tag_e,
                    i_s + len(p_tag_e) + len(p_tag_s) + (len(p_str) - slen),
                )  # search after, start + tags sizes + diff in inserted text.
                applyError = True
                s.add_fault("incomplete tag", "incomplete " + p_tag_s + "")
            #
            return p_str

        ###

        # p_obj.text = search_str(p_obj.text, s.TAG_S, s.TAG_E)
        # replacing break images and style.
        search_str(p_obj.text, s.TAG_S, s.TAG_E)

        if applyError:
            # plpy.notice("Settings color - {}".format(p_objpar.text))
            if p_objpar.text.find(s.STR_ERR) < 0:
                r = p_objpar.add_run()
                r.text = s.STR_ERR + " " + errorstr + ".  "
                font = r.font
                font.color.rgb = RGBColor(0xFF, 0x01, 0x01)
            #
        #

    #

    def mark_docx(s):
        try:

            if len(s.inputdata) <= 0:
                raise MergeToolWarning("No tag data given.")
            #

            errortags = s.inputdata.get("marktags")
            if errortags is None:
                raise Exception(
                    "Expected an array list of tags. Got: {}".format(s.inputdata)
                )
            #

            s.load_docx()

            for p in s.srcdoc.paragraphs:
                if len(p.runs) > 0:
                    for r in p.runs:
                        s.markerr_obj(r, p, errortags)
                    #
                else:
                    s.markerr_obj(p, p, errortags)
                #
            #

            global tblcnt

            def looptbl(p_tbls, p_parent=0):
                global tblcnt
                if p_parent == 0:
                    tblcnt = 0

                for t in p_tbls:
                    tblcnt += 1
                    for r in t.rows:
                        for c in r.cells:
                            for p in c.paragraphs:
                                if len(p.runs) > 0:
                                    for r in p.runs:
                                        s.markerr_obj(
                                            r,
                                            p,
                                            errortags,
                                            "{},{}".format(tblcnt, p_parent),
                                        )
                                    #
                                else:
                                    s.markerr_obj(
                                        p,
                                        p,
                                        errortags,
                                        "{},{}".format(tblcnt, p_parent),
                                    )
                                #
                            #
                            if len(c.tables) > 0:
                                looptbl(c.tables, tblcnt)
                            #
                        #
                    #
                #

            #
            looptbl(s.srcdoc.tables)
            s.save_docx()

        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            raise MergeToolError("Failed marking docx: " + str(e), "mark_docx", exc_tb)
        #

    #

    def convert_docx(s, p_mark=True):
        replacekeys = {}
        eval_tag_start = 0
        eval_tag_end = 0

        try:
            tagdata = s.inputdata.get("tags")
            if s.inputdata is None or len(s.inputdata) <= 0:
                raise MergeToolWarning("No tag data given.")
            #
            replacekeys = s.inputdata.get(
                "replacetags"
            )  ##expected {"replacetags": {"key1":"newkey1", "key2":"newkey2"}
            s.load_docx()

            for p in s.srcdoc.paragraphs:
                if len(p.runs) > 0:
                    for r in p.runs:
                        if replacekeys is not None:
                            if (
                                r.text.find(s.TAG_S) >= 0
                                or r.text.find(s.TAG_S_OLD) >= 0
                                or r.text.find(s.TAG_S_SP) >= 0
                            ):
                                r.text = s.replace_str(r.text, replacekeys)
                            #
                        #
                        if p_mark:
                            s.markerr_obj(r, p)
                        #

                    #
                else:
                    if replacekeys is not None:
                        if (
                            p.text.find(s.TAG_S) >= 0
                            or p.text.find(s.TAG_S_OLD) >= 0
                            or p.text.find(s.TAG_S_SP) >= 0
                        ):
                            p.text = s.replace_str(p.text, replacekeys)
                        #
                    #
                    if p_mark:
                        s.markerr_obj(p, p)
                    #
                #
            #
            for t in s.srcdoc.tables:
                for r in t.rows:
                    for c in r.cells:
                        for p in c.paragraphs:
                            if len(p.runs) > 0:
                                for r in p.runs:
                                    if (
                                        replacekeys is not None
                                        and (
                                            r.text.find(s.TAG_S) >= 0
                                            or r.text.find(s.TAG_S_OLD) >= 0
                                        )
                                        or r.text.find(s.TAG_S_SP) >= 0
                                    ):
                                        r.text = s.replace_str(r.text, replacekeys)
                                    #
                                    if p_mark:
                                        s.markerr_obj(r, p)
                                    #
                                #
                            else:
                                if (
                                    replacekeys is not None
                                    and (
                                        p.text.find(s.TAG_S) >= 0
                                        or p.text.find(s.TAG_S_OLD) >= 0
                                    )
                                    or p.text.find(s.TAG_S_SP) >= 0
                                ):
                                    p.text = s.replace_str(p.text, replacekeys)
                                #
                                if p_mark:
                                    s.markerr_obj(p, p)
                                #
                            #
                        #
                    #
                #
            #

            s.save_docx()
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            raise MergeToolError(
                "Failed converting docx: " + str(e), "convert_docx", exc_tb
            )
        #

    #

    def convertonly_docx(s, oldtag_s, oldtag_e):
        replacekeys = {}

        try:
            s.load_docx()
            # plpy.notice("Converting Document: {} from tags: {},{} to {},{} and saved to {}".format(s.srcfilename, oldtag_s, oldtag_e,s.TAG_S, s.TAG_E, s.destfilename))
            for p in s.srcdoc.paragraphs:
                if len(p.runs) > 0:
                    for r in p.runs:
                        if r.text.find(oldtag_s) >= 0 or r.text.find(oldtag_s) >= 0:
                            r.text = r.text.replace(oldtag_s, s.TAG_S)
                            r.text = r.text.replace(oldtag_e, s.TAG_E)
                        #
                    #
                else:
                    if p.text.find(oldtag_s) >= 0 or p.text.find(oldtag_s) >= 0:
                        p.text = p.text.replace(oldtag_s, s.TAG_S)
                        p.text = p.text.replace(oldtag_e, s.TAG_E)

                    #
                #
            #
            def looptbl(t):
                for r in t.rows:
                    for c in r.cells:
                        for p in c.paragraphs:
                            if len(c.tables) > 0:
                                for t2 in c.tables:
                                    looptbl(t2)
                                #
                            #
                            if len(p.runs) > 0:
                                for r in p.runs:
                                    if (
                                        r.text.find(oldtag_s) >= 0
                                        and r.text.find(oldtag_s) >= 0
                                    ):
                                        r.text = r.text.replace(oldtag_s, s.TAG_S)
                                        r.text = r.text.replace(oldtag_e, s.TAG_E)
                                    #
                                #
                            else:
                                if (
                                    p.text.find(oldtag_s) >= 0
                                    and p.text.find(oldtag_s) >= 0
                                ):
                                    p.text = p.text.replace(oldtag_s, s.TAG_S)
                                    p.text = p.text.replace(oldtag_e, s.TAG_E)
                                #
                            #
                        #
                    #
                #

            #

            for t in s.srcdoc.tables:
                looptbl(t)
            #

            s.actionHeadFootTags("head", "convert")
            s.actionHeadFootTags("foot", "convert")

            s.save_docx()
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            raise MergeToolError(
                "Failed converting docx: " + str(e), "convertonly_docx", exc_tb
            )
        #

    #

    def setTableName(s, p_tblparentid, p_str, p_start=0):
        p_tag_s = s.TAG_S_POSTOP
        p_tag_e = s.TAG_E_POSTOP

        i_s = p_str.find(p_tag_s, p_start)
        i_e = p_str.find(p_tag_e, i_s)
        if i_s >= 0 and i_e > i_s:
            keyname = p_str[i_s + len(p_tag_s) : i_e]
            if len(keyname) > 1:
                keyparts = str(keyname).split(":")
                if (
                    keyparts[0] == "name"
                    and s.tbl_map.get(str(p_tblparentid), None) is not None
                ):
                    s.tbl_map[str(p_tblparentid)]["name"] = keyparts[1]
                    return keyparts[1]
                #
            #
            # search next tag
            res = s.setTableName(p_tblparentid, p_str, i_e + len(p_tag_e))
            if res != "":
                return res
            #
        #
        return

    #

    def tbl_basic_replace(s, p_tables, p_parent=0):
        if p_parent == 0:
            s.tbl_basic_replace_cnt = 0

        for t in p_tables:
            try:
                s.tbl_basic_replace_cnt += 1
                tblsource = "{},{}".format(s.tbl_basic_replace_cnt, p_parent)
                # plpy.notice("Table {} Parent: {} -> {}".format(tblsource,p_parent, t))
                if p_parent > 0:
                    if s.tbl_map.get(str(p_parent)) is None:
                        s.tbl_map[str(p_parent)] = {
                            "tblptr": t,
                            "children": None,
                            "nr": tblsource,
                            "limitrow": 0,
                            "name": "",
                        }
                        # raise Exception("New Parent: {} No parent: {}, Tables:{}".format(p_parent,  s.tbl_map,len(p_tables)))
                    if s.tbl_map[str(p_parent)]["children"] is None:
                        s.tbl_map[str(p_parent)]["children"] = {}
                    s.tbl_map[str(p_parent)]["children"][
                        str(s.tbl_basic_replace_cnt)
                    ] = {
                        "tblptr": t,
                        "children": None,
                        "nr": tblsource,
                        "limitrow": 0,
                        "name": "",
                    }
                else:
                    s.tbl_map[str(s.tbl_basic_replace_cnt)] = {
                        "tblptr": t,
                        "children": None,
                        "nr": tblsource,
                        "limitrow": 0,
                        "name": "",
                    }
                #
                lastid = s.tbl_basic_replace_cnt
                distinctdata = []
                rowcnt = -1
                for row in t.rows:
                    colidx = 0
                    rowcnt += 1
                    rmrow = False
                    for cell in row.cells:
                        s.setTableName(s.tbl_basic_replace_cnt, cell.text)
                        rmcol = False
                        for p in cell.paragraphs:
                            if len(p.runs) > 0:
                                for rn in p.runs:
                                    opdata = s.tag_value_replace(
                                        rn,
                                        p,
                                        None,
                                        None,
                                        "tablebasic",
                                        str(lastid),
                                        colidx,
                                    )
                                    if opdata is None:
                                        continue
                                    if rn.text.find(s.STR_TAG_FREETBL) > 0:
                                        s.tbl_mode[str(p_parent)] = "free"
                                    # remove the blank value row again. Note that its not the filter and depends on value of field with this remove tag.

                                    # sqlmsg("limitrow opdata {}".format(opdata))
                                    if tryInt(opdata.get("rmblnk", "0")) == 2:
                                        rmrow = True
                                    if tryInt(opdata.get("rmblnkcol", "0")) == 2:
                                        rmcol = True
                                    if tryInt(opdata.get("rmdr", "0")) == 2:
                                        rmrow = True
                                    if tryInt(opdata.get("limitrow", "0")) > 0:
                                        # sqlmsg("limitrow 2 opdata {}".format(opdata))
                                        if (
                                            s.tbl_map.get(str(s.tbl_basic_replace_cnt))
                                            is None
                                        ):
                                            s.tbl_map[str(s.tbl_basic_replace_cnt)] = {}
                                        #
                                        s.tbl_map[str(s.tbl_basic_replace_cnt)][
                                            "limitrow"
                                        ] = opdata.get("limitrow", "0")
                                    #

                                    if (
                                        len(opdata.get("mergedtags", [])) > 0
                                    ):  # this will reset the tag, important to detect sibling
                                        rmrow = False
                                        rmcol = False
                                    #
                                    s.tag_picture_add(rn)
                                #
                            else:
                                opdata = s.tag_value_replace(
                                    p,
                                    None,
                                    None,
                                    None,
                                    "tablebasic",
                                    str(lastid),
                                    colidx,
                                )
                                if p.text.find(s.STR_TAG_FREETBL) > 0:
                                    s.tbl_mode[str(p_parent)] = "free"

                                # remove the blank value row again. Note that its not the filter and depends on value of field with this remove tag.
                                if opdata.get("rmblnk", 0) == 2:
                                    rmrow = True
                                if opdata.get("rmblnkcol", 0) == 2:
                                    rmcol = True
                                if opdata.get("rmdr", 0) == 2:
                                    rmrow = True
                                if opdata.get("limitrow", 0) > 0:
                                    s.tbl_map[str(s.tbl_basic_replace_cnt)][
                                        "limitrow"
                                    ] = opdata["limitrow"]
                                #

                                if (
                                    len(opdata.get("mergedtags", [])) > 0
                                ):  # this will reset the tag, important to detect sibling
                                    rmrow = False
                                    rmcol = False
                                #
                            #
                        #
                        if len(cell.tables) > 0:
                            # plpy.notice("Inner Table {} Parent: {} Cell: {} Row: {}".format(tblsource,p_parent, colidx,rowidx))
                            s.tbl_basic_replace(cell.tables, lastid)
                        #

                        if rmcol:
                            s.tbl_rmcol.append(cell)
                        #
                        colidx += 1
                    #
                    if rmrow:
                        s.tbl_rmrow.append(row)
                    #
                #
            except Exception as e:
                tb = traceback.format_exc()

                sqlmsg(
                    "Table Replace Error: {} Trace: {}".format(e, tb),
                    p_type="local error",
                )
            #
            #
        #

    #

    def getRunNearestParentTable(s, p_obj):
        if isinstance(p_obj, docx.text.run.Run):
            par = p_obj._parent
            if isinstance(par, docx.table._Cell):
                tbl = par._parent
                return tbl
            #
            gran = par._parent
            if isinstance(gran, docx.table._Cell):
                tbl = gran._parent
                return tbl
            #
        #
        return None

    #

    def tbl_post_loop(s, p_tables, p_parent=0):
        if p_parent == 0:
            s.tbl_basic_replace_cnt = 0

        # plpy.notice("s.tbl_complex_rowdat={}".format(s.tbl_complex_rowdat))

        for r in s.tbl_rmcol:
            # sqlmsg("Removing cell: len: {} -> {}".format(len(s.tbl_rmcol), r) )
            try:
                parent = r._tc.find("..")
                if parent is not None:
                    parent.remove(r._tc)
                    # sqlmsg("Removed col: {}".format(r))
                #
            except Exception as e:
                sqlmsg("Removing cell failed -> {}".format(e))
            #
        #

        rowremoveList = []

        for tbl in p_tables:
            s.tbl_basic_replace_cnt += 1
            tblsource = "{},{}".format(s.tbl_basic_replace_cnt, p_parent)
            lastid = s.tbl_basic_replace_cnt
            # plpy.notice("t={}".format(lastid))

            rmRowIndicator = {}
            rowcnt = 0
            ## Tag rows delete
            for row in tbl.rows:
                colidx = 0
                rowcnt += 1
                for cell in row.cells:
                    if cell.text.find(s.TAG_COND_RTBL) >= 0:
                        # plpy.notice("Remove table test: {}, {}".format(lastid,s.tbl_complex_rowdat.get(str(lastid))))
                        if s.tbl_complex_rowdat.get(str(lastid), {}).get("max", 1) <= 0:
                            parent = tbl._tbl.find("..")
                            if parent is not None:
                                parent.remove(tbl._tbl)
                                break
                            #
                        #
                    #
                    # sqlmsg("Cell Text: {} , {}".format(rowcnt, cell.text))
                    if cell.text.find(s.TAG_COND_RMR) >= 0:
                        # plpy.notice("Removing Row: {}".format(rowcnt))
                        # sqlmsg("Removing Row: {} , {}".format(rowcnt, cell.text))
                        rmRowIndicator["{}".format(rowcnt)] = True
                        parent = row._tr.find("..")
                        if parent is not None:
                            # parent.remove(row._tr)
                            rowremoveList.append(
                                {"row": row, "parent": parent, "rowcnt": rowcnt}
                            )
                            continue
                        #
                    #
                    if cell.text.find(s.TAG_COND_CLR) >= 0:
                        cell.text = ""
                        continue
                    #
                    if cell.text.find(s.TAG_COND_RMC) >= 0:
                        # plpy.notice("Remove Col {}".format(colidx))
                        for c2 in tbl.columns[colidx].cells:
                            # plpy.notice("Remove Cell: {}".format(c2.text))
                            c2._tc.clear_content()
                            # parent = c2._tc.find('..')
                            # if parent is not None:
                            #   parent.remove(c2._tc)
                            # #
                        #
                        continue
                    #

                    # for par in cell.paragraphs:
                    #   for r2 in par.runs:
                    #     if len(r2.text) < 2:
                    #       sqlmsg("__ Tbl: {} run: {} par:{}".format(tblsource,r2.text,par.text ))
                    #     #
                    #     if str(r2.text) == " " or str(r2.text) == ".":
                    #       sqlmsg("__Clear Tbl: {} run: {}".format(tblsource,r2), p_type="local error")
                    #       parent = r2._r.find('..')
                    #       if r2._r is not None:
                    #         parent.remove(r2._r)
                    #       #
                    #     #
                    #   #
                    # #
                    colidx += 1
                #
            #

            ##merge missing and special tags
            rowcnt = 0
            totalrows = len(tbl.rows)
            for row in tbl.rows:
                colidx = -1
                rowcnt += 1
                for cell in row.cells:
                    colidx += 1
                    for par in cell.paragraphs:
                        for r2 in par.runs:
                            tags = s.peek_str(r2.text, False)
                            s.peek_alias_str(
                                r2.text, tblsource, "peekdeclare", r2, rowcnt
                            )
                            # if len(tags) <= 0: continue
                            # plpy.notice("Left over tags found {} ".format(tags))

                            for tag in tags:
                                tagname = ""
                                tagdat = ""
                                tagop = ""
                                ichr = tag.find(s.TAG_SPLT)
                                if ichr >= 0:
                                    tagname = tag[:ichr].replace(s.TAG_S, "")
                                    tagdat = tag[ichr + 1 :].replace(s.TAG_E, "")
                                    # sqlmsg("Chk Kickstrip {} tag: {} tagname: {} tagdat: {} tagop:{}".format(rowcnt,tag, tagname,tagdat,tagop))
                                else:
                                    if tagname == "":
                                        tagname = tag.replace(s.TAG_S, "").replace(
                                            s.TAG_E, ""
                                        )
                                    #
                                #

                                ichr = tag.find(s.TAG_RULESPLT)
                                if ichr >= 0:
                                    tagname = tag[:ichr].replace(s.TAG_S, "")
                                    tagop = tag[ichr + 1 :].replace(s.TAG_E, "")
                                else:
                                    if tagname == "":
                                        tagname = tag.replace(s.TAG_S, "").replace(
                                            s.TAG_E, ""
                                        )
                                    #
                                #

                                if tagname == "wordarticle":
                                    s.globalTagsReplace(r2)
                                    continue
                                #
                                # sqlmsg("Chk ss {} tag: {} tagname: {} tagdat: {} tagop:{}".format(rowcnt,tag, tagname,tagdat,tagop))

                                if tagname.find(s.STR_TAG_ESIGN) >= 0:
                                    continue
                                #

                                if rmRowIndicator.get("{}".format(rowcnt), False):
                                    continue
                                #

                                # sqlmsg("Chk tagname: {} tagdat: {}".format(tagname, tagop))
                                def opReplace(pStr, pTag, pVal, pFullTag=None):
                                    try:
                                        if tagop != "":
                                            rv = plpy.execute(
                                                """
                        select r.result from exec(format('select (%s)::citext',mailmerge_rule($S$'{0}'$S$, '{2}', '{0}', '{1}') ))  r (result citext)
                      """.format(
                                                    pVal, pTag, tagop
                                                )
                                            )
                                            pVal = rv[0]["result"]
                                        #
                                    except Exception as e:
                                        sqlmsg(
                                            "Failed to replace operator value for oridinals: tag: {} op: {} Error: {}".format(
                                                tagname, tagop, e
                                            ),
                                            p_type="local notice",
                                        )
                                    #
                                    # sqlmsg("Chk opReplace pTag: {} pVal: {} pFullTag: {} pStr:{}".format(pTag, pVal,pFullTag,pStr))
                                    return pStr.replace(pTag, pVal)

                                #

                                if (
                                    tagname in s.OP_TAGS
                                ):  ##Handle the ordinal/cardinal values
                                    if s.tbl_ordinaldata.get(tblsource) is None:
                                        s.tbl_ordinaldata[tblsource] = {}
                                    #
                                    if (
                                        s.tbl_ordinaldata.get(tblsource).get(tagname)
                                        is None
                                    ):
                                        if tagdat.strip(" ") == "":
                                            tagdat = "1"
                                        s.tbl_ordinaldata[tblsource][tagname] = tryInt(
                                            tagdat
                                        )
                                    #
                                    rmlist = 0

                                    for elem in s.table_removerow_list:
                                        if elem is None:
                                            continue
                                        if elem.get("xmlptr", None) is None:
                                            continue
                                        # sqlmsg("__elem {} nearTable {}".format(elem["xmlptr"], tbl))
                                        if elem.get("tableid", "") == tblsource:
                                            rmlist = rmlist + 1
                                            # sqlmsg("Found elem {} rmlist {}".format(elem, rmlist))
                                        #
                                    #

                                    if rmlist > 0:
                                        s.tbl_ordinaldata[tblsource][tagname] = (
                                            rowcnt - rmlist
                                        )
                                    #

                                    # sqlmsg("OPS - Tag:{} Ordinal:{}, Dat:{} rmlist:{} DATA:{}".format(tagname, s.tbl_ordinaldata[tblsource][tagname], tagdat,rmlist, s.tbl_ordinaldata))
                                    # plpy.notice("Handle cardinal/ordinals for {},{} tblsrc={} ordat={} tagname={} rowcnt={}".format(tag,tagname, tblsource, s.tbl_ordinaldata[tblsource].get(tagname), tagname, rowcnt))

                                    if tagname == "ordinal":
                                        r2.text = opReplace(
                                            r2.text,
                                            tag,
                                            str(
                                                ordinalstr(
                                                    s.tbl_ordinaldata[tblsource][
                                                        tagname
                                                    ]
                                                )
                                            ),
                                        )
                                    elif tagname == "ordinalw":
                                        r2.text = opReplace(
                                            r2.text,
                                            tag,
                                            n2w(
                                                s.tbl_ordinaldata[tblsource][tagname],
                                                p_type="ordinal",
                                            ),
                                        )
                                    elif tagname == "cardinal":
                                        r2.text = opReplace(
                                            r2.text,
                                            tag,
                                            str(s.tbl_ordinaldata[tblsource][tagname]),
                                        )
                                    elif tagname == "cardinalw":
                                        r2.text = opReplace(
                                            r2.text,
                                            tag,
                                            n2w(
                                                s.tbl_ordinaldata[tblsource][tagname],
                                                p_type="cardinal",
                                            ),
                                        )
                                    #

                                    s.tbl_ordinaldata[tblsource][tagname] += 1
                                elif tagname == "minval":
                                    r2.text = opReplace(
                                        r2.text,
                                        tag,
                                        str(
                                            s.tbl_complex_rowdat.get(
                                                str(lastid), {}
                                            ).get("min", 1)
                                        ),
                                    )
                                elif tagname == "minvalw":
                                    r2.text = opReplace(
                                        r2.text,
                                        tag,
                                        n2w(
                                            s.tbl_complex_rowdat.get(
                                                str(lastid), {}
                                            ).get("min", 1)
                                        ),
                                    )
                                elif tagname == "maxval":
                                    r2.text = opReplace(
                                        r2.text,
                                        tag,
                                        str(
                                            s.tbl_complex_rowdat.get(
                                                str(lastid), {}
                                            ).get("max", "")
                                        ),
                                    )
                                elif tagname == "maxvalw":
                                    r2.text = opReplace(
                                        r2.text,
                                        tag,
                                        n2w(
                                            s.tbl_complex_rowdat.get(
                                                str(lastid), {}
                                            ).get("max", "")
                                        ),
                                    )
                                #
                                elif tagname == "..tblsrc":
                                    r2.text = opReplace(r2.text, tag, tblsource)
                                #
                                elif tagname == "..lastid":
                                    r2.text = opReplace(r2.text, tag, str(lastid))
                                # elif tagop != "":
                                #   sqlmsg("Unhandled post merge operator : tag: {} op: {}".format(tagname, tagop), p_type="local notice")
                                # #

                                else:
                                    if (
                                        s.cleanMissing
                                        and not tag.lower().find(s.STR_TAG_ESIGN) >= 0
                                    ):
                                        r2.text = r2.text.replace(tag, "")
                                        # sqlmsg("Leftover tag empty {} ".format(tag))
                                        # pass
                                #
                            #
                            s.eval_replace(r2, "run", tblsource)

                            # for tobj in s.table_removerow_list:
                            #   if tobj.get("xmlptr") == tbl:
                            #     removedtotal = totalrows - tryInt(tobj.get("total",0) - tobj.get("row",0))
                            #     sqlmsg("Found XML {}={} T={} len={} {},{}".format(rowcnt, tobj.get("row",0), removedtotal ,len(s.table_removerow_list), tobj, tbl))

                            #     if removedtotal < rowcnt:
                            #       skipcnt = True
                            #       sqlmsg("Skip Cnt XML {},{},{} : {}".format(rowcnt,removedtotal,totalrows,tobj.get("row",0) ))
                            #     #
                            #   #
                            # #
                            #
                            #

                            if r2.text.find(s.TAG_COND_RMC) >= 0:
                                r2.text = r2.text.replace(s.TAG_COND_RMC, "")
                            #
                            if r2.text.find(s.TAG_COND_RMR) >= 0:
                                r2.text = r2.text.replace(s.TAG_COND_RMR, "")
                            #
                            if r2.text.find(s.TAG_COND_RTBL) >= 0:
                                r2.text = r2.text.replace(s.TAG_COND_RTBL, "")
                            #
                            if r2.text.find(s.TAG_COND_RMP) >= 0:
                                r2.text = r2.text.replace(s.TAG_COND_RMP, "")
                                par.clear()
                            #
                            s.tag_picture_add(r2)
                            skipremoved = False
                            for i in rowremoveList:
                                if i["rowcnt"] == rowcnt:
                                    skipremoved = True
                                #
                            #
                            if not skipremoved:
                                s.col_operation_compute(
                                    r2, tblsource, colidx, rowcnt, row
                                )
                                s.eval_replace(r2, "run", tblsource)
                            #
                            # sqlmsg("s.tbl_search for {} and is {}".format(tblsource, s.tbl_search))
                            if s.tbl_search.get(str(lastid)) is not None:
                                searchvars = s.tbl_search.get(str(lastid), {})
                                if searchvars.get("_ptr") is None:
                                    searchvars["_ptr"] = tbl
                                #

                                if (
                                    par.text.lower().find(
                                        searchvars.get("searchvalue").lower()
                                    )
                                    >= 0
                                ):
                                    searchvars["found"] = True
                                #
                                sqlmsg(
                                    "@@tbl_search for {} and is {}".format(
                                        tblsource, s.tbl_search
                                    )
                                )
                            #
                    #

                    #
                    if len(cell.tables) > 0:
                        s.tbl_post_loop(cell.tables, lastid)
                    #

                #
            #

        #

        # sqlmsg("Removing Row Data: {}".format(rowremoveList))
        for i in rowremoveList:
            # sqlmsg("Removing Row: {}".format(i))
            if i["parent"] is not None:
                try:
                    i["parent"].remove(i["row"]._tr)
                except Exception as e:
                    sqlmsg("Removing Row Failed: {}".format(e))
                #
            #
        #

        for i in s.tbl_search:
            if (
                s.tbl_search[i].get("_ptr") is not None
                and s.tbl_search[i].get("action", "none") == "remove"
                and s.tbl_search[i].get("found", True) == False
            ):
                tblptr = s.tbl_search[i].get("_ptr")
                parent = tblptr._tbl.find("..")
                if parent is not None:
                    parent.remove(tblptr._tbl)
                #
                # sqlmsg("s.tbl_search Remove {}, {}".format(i,s.tbl_search[i]))
                # parent = row._tr.find('..')
            #
        #

    #

    def col_operation_search_str(s, p_str, p_start=0):
        p_tag_s = s.TAG_S_POSTOP
        p_tag_e = s.TAG_E_POSTOP
        tags = []

        i_s = p_str.find(p_tag_s, p_start)
        i_e = p_str.find(p_tag_e, i_s)
        if i_s >= 0 and i_e > i_s:
            keyname = p_str[i_s + len(p_tag_s) : i_e]
            if len(keyname) > 1:
                keyparts = str(keyname).split(":")
                cellindex = None

                type = ""
                if keyparts[0] in s.CELL_OP_TAGS:
                    type = "cell"
                    if len(keyparts) > 1:
                        cellindex = tryInt(keyparts[1])
                    #
                #
                if keyparts[0] in ["debug"]:
                    type = "debug"
                #
                if keyparts[0] in ["ref"]:
                    type = "ref"
                #
                if keyparts[0] in ["name"]:
                    type = "name"
                #

                if type == "":
                    return tags
                tags.append(
                    {
                        "tagfull": "{}{}{}".format(p_tag_s, keyname, p_tag_e),
                        "tag": keyparts[0],
                        "operators": keyparts[1:],
                        "cellindex": cellindex,
                        "type": type,
                    }
                )
                # recursive to next position.
                moretags = s.col_operation_search_str(p_str, i_e + len(p_tag_e))
                tags.extend(moretags)
            #
        #
        return tags

    #

    def col_operation_compute(s, p_obj, p_tblid, p_colidx, p_rowidx, p_rowptr):
        ops = s.tbl_col_operation.get("{0}".format(p_tblid), None)
        if ops is None:
            # sqlmsg("Col_operation_compute nulldata -> index:{}  colopts:{}".format("{0}_{1}".format(p_tblid,p_colidx), s.tbl_col_operation))
            return
        #
        for opdata in ops:
            if opdata is None:
                continue
            if getProp(opdata, "skip", "") != "":
                continue
            tagfull = getProp(opdata, "tagfull", "")
            cellindex = getProp(opdata, "cellindex", -1)
            rowptr = getProp(opdata, "rowptr", None)
            if getProp(opdata, "type", "") != "cell":
                continue
            #
            if tagfull == "":
                # sqlmsg("Col_operation_compute notag -> opdata:{} index:{}  colopts:{}".format(opdata,"{0}_{1}".format(p_tblid,p_colidx), s.tbl_col_operation))
                continue
            #
            if p_obj is None:
                continue
            isfinal = p_obj.text.find(tagfull) >= 0
            issource = False
            if rowptr is not None:
                issource = rowptr._tr == p_rowptr._tr
            #
            skipRow = False
            # remove the already removed rows
            for i, v in enumerate(s.table_removerow_list):
                if v["rowptr"] is not None and v["rowptr"]._tr == p_rowptr._tr:
                    skipRow = True
                    break
                #
            #

            if skipRow:
                continue
            #

            try:
                if isfinal:
                    # sqlmsg("Col_operation_compute debug -> opdata:{} issource:{} p_colidx:{} cellindex:{} tbl_col_operation:{}".format(opdata,issource,p_colidx,cellindex,s.tbl_col_operation ))

                    if len(opdata["values"]) == 0:
                        p_obj.text = p_obj.text.replace(tagfull, "")
                        opdata["result"] = None
                    #

                    if tagfull.find("max") >= 0:
                        maxval = round(max_any(opdata["values"]), 2)
                        opdata["result"] = "{:.2f}".format(maxval)
                        p_obj.text = p_obj.text.replace(
                            tagfull, "{:.2f}".format(maxval)
                        )
                    #

                    if tagfull.find("min") >= 0:
                        minval = round(min_any(opdata["values"]), 2)
                        opdata["result"] = "{:.2f}".format(minval)
                        p_obj.text = p_obj.text.replace(
                            tagfull, "{:.2f}".format(minval)
                        )
                    #

                    if tagfull.find("avg") >= 0:
                        avgval = round(avg_any(opdata["values"]), 2)
                        opdata["result"] = "{:.2f}".format(avgval)
                        p_obj.text = p_obj.text.replace(
                            tagfull, "{:.2f}".format(avgval)
                        )
                    #

                    if tagfull.find("sum") >= 0:
                        sumval = round(sum_any(opdata["values"]), 2)
                        opdata["result"] = "{:.2f}".format(sumval)
                        p_obj.text = p_obj.text.replace(
                            tagfull, "{:.2f}".format(sumval)
                        )
                    #

                    if tagfull.find("count") >= 0:
                        sumval = len(opdata["values"])
                        opdata["result"] = sumval
                        p_obj.text = p_obj.text.replace(
                            tagfull, "{:.2f}".format(sumval)
                        )
                    #

                    s.eval_replace(p_obj, p_type="docx", p_tableid=p_tblid)

                elif cellindex == p_colidx and not issource:
                    # sqlmsg("Col_operation_compute debug -> p_tblid:{} row/col:{}/{} colopts:{}".format(p_tblid,p_rowidx,p_colidx, opdata))

                    numval = getNumber(p_obj.text, None)
                    if numval is not None:
                        opdata["values"].append(numval)
                    elif (
                        str(p_obj.text).strip() != ""
                        and p_obj.text.find(s.TAG_S_POSTOP) < 0
                    ):
                        opdata["values"].append(p_obj.text)
                        opdata["mixedtypes"] = True
                    #
                    opdata["rows"] = len(opdata["values"])
                    if p_rowidx > opdata["rows"]:
                        opdata["rows"] = p_rowidx
                    #
                #
            except Exception as e:
                sqlmsg(
                    "Col_operation_compute failed -> {} opdata: {}".format(e, opdata)
                )
            #
        #

    #

    def basic_operation_compute(s, p_obj, p_location="_base", p_overrideDetail=None):
        tagdataList = s.col_operation_search_str(p_obj.text)
        for tagdata in tagdataList:
            if tagdata is None:
                continue
            if s.tbl_col_operation.get(p_location, None) is None:
                s.tbl_col_operation[p_location] = []
            #
            obj = {
                "tag": getProp(tagdata, "tag", ""),
                "tagfull": getProp(tagdata, "tagfull", ""),
                "type": getProp(tagdata, "type", ""),
                "cellindex": getProp(tagdata, "cellindex", -1),
                "rowidx": getProp(tagdata, "rowidx", -1),
                "values": [],
                "result": None,
                "rowptr": None,
                "operators": getProp(tagdata, "operators", []),
            }
            if p_overrideDetail is not None:
                obj["cellindex"] = getProp(
                    p_overrideDetail, "cellindex", obj["cellindex"]
                )
                obj["rowidx"] = getProp(p_overrideDetail, "rowidx", obj["rowidx"])
                obj["rowptr"] = getProp(p_overrideDetail, "rowptr", obj["rowptr"])
            #
            if getProp(tagdata, "cellindex", -1) >= 0:
                obj["cellindex"] = tagdata.get("cellindex", obj["cellindex"])
            #
            if obj is not None:
                s.tbl_col_operation[p_location].append(obj)

                if tagdata["type"] == "debug" and len(tagdata["operators"]) == 0:
                    p_obj.text = p_obj.text.replace(
                        tagdata["tagfull"],
                        "{}".format(
                            json.dumps(s.tbl_col_operation, cls=BJSONEnc, indent=4)
                        ),
                    )
                elif (
                    tagdata["type"] == "debug" and tagdata["operators"][0] == "tablemap"
                ):
                    p_obj.text = p_obj.text.replace(
                        tagdata["tagfull"],
                        "{}".format(json.dumps(s.tbl_map, cls=BJSONEnc, indent=4)),
                    )
                elif tagdata["type"] == "debug" and tagdata["operators"][0] == "json":
                    p_obj.text = p_obj.text.replace(
                        tagdata["tagfull"],
                        "{}".format(json.dumps(s.inputdata, cls=BJSONEnc, indent=4)),
                    )
                elif tagdata["type"] == "debug" and tagdata["operators"][0] == "peek":
                    p_obj.text = p_obj.text.replace(
                        tagdata["tagfull"],
                        "{}".format(json.dumps(s.peek_docx(), cls=BJSONEnc, indent=4)),
                    )
                #

                if tagdata["type"] == "ref" and len(tagdata["operators"]) > 2:
                    tableid = tagdata["operators"][0]
                    type = tagdata["operators"][1]
                    # type = tagdata["operators"][3]
                    reflist = s.tbl_col_operation.get(tableid, [])
                    if len(reflist) == 0:
                        for t in s.tbl_map:
                            o = s.tbl_map[t]
                            if o.get("name", "").lower() == tableid.lower():
                                reflist = s.tbl_col_operation.get(o.get("nr", ""), [])
                                break
                            #
                        #
                    #
                    for o in reflist:
                        if o.get("tag", "").find(type) >= 0:
                            p_obj.text = p_obj.text.replace(
                                tagdata["tagfull"], "{}".format(o.get("result", ""))
                            )
                        #
                    #
                #
                if tagdata["type"] == "name":
                    p_obj.text = p_obj.text.replace(tagdata["tagfull"], "")
                #
                s.eval_replace(p_obj, p_type="docx", p_tableid=p_location)
            #
        #

    #

    def tbl_complex_headerstart(s, p_tbl, p_taglist, p_tblid, p_fields):
        rownr = 0
        templaterows = {}
        tblsource = s.tbl_map.get(p_tblid, {}).get("nr", "")
        if p_tbl is None:
            return templaterows
        # if not hasattr(p_tbl,"rows"): return templaterows
        if s.tbl_complex_rowdat.get(p_tblid) is None:
            s.tbl_complex_rowdat[p_tblid] = {"min": 0, "max": 0}
        #
        # process col operations

        rowindex = 0
        for row in p_tbl.rows:
            rowindex += 1
            cellindex = 0
            skip = False
            for cell in row.cells:
                if cell.text.find("#rmr") >= 0 or cell.text.find(s.TAG_COND_RMR) >= 0:
                    skip = True
                    break
                #
            #
            if skip:
                continue

            for cell in row.cells:
                s.basic_operation_compute(
                    cell,
                    tblsource,
                    {
                        "cellindex": cellindex,
                        "rowidx": rowindex,
                        "rowptr": row,
                        "cellptr": cell,
                    },
                )
                #
                cellindex += 1
            #
        #

        for row in p_tbl.rows:
            breakrow = False
            for cell in row.cells:
                bodytext = cell.text

                # check fields in paragraph
                for fieldname in p_taglist:
                    if type(fieldname) is dict:
                        continue  # we dont want the sub table data here.
                    if bodytext.find(fieldname) >= 0:
                        # if int(p_fields.get(fieldname,{}).get("type", 0)) == 4:
                        #  continue
                        #
                        templaterows[rownr] = {
                            "row": row,
                            "rowtrcopy": copy.deepcopy(row._tr),
                            "merged": False,
                            "type": int(p_fields.get(fieldname, {}).get("type", 0)),
                        }
                        rownr += 1
                        breakrow = True
                        if s.tbl_complex_rowdat[p_tblid]["min"] <= 0:
                            s.tbl_complex_rowdat[p_tblid]["min"] = rownr
                        break
                    #
                    # handle blank row
                    elif bodytext.find(s.TAG_S + "blnkrow" + s.TAG_E) >= 0:
                        templaterows[rownr] = {
                            "row": row,
                            "rowtrcopy": copy.deepcopy(row._tr),
                            "merged": False,
                            "type": 0,
                        }
                        rownr += 1
                        breakrow = True
                        if s.tbl_complex_rowdat[p_tblid]["min"] <= 0:
                            s.tbl_complex_rowdat[p_tblid]["min"] = rownr
                        break
                    #
                #
                if breakrow:
                    break
                for childtable in cell.tables:
                    for childrow in childtable.rows:
                        for childcell in childrow.cells:
                            bodytext = childcell.text
                            for fieldname in p_taglist:
                                if type(fieldname) is dict:
                                    continue  # we dont want the sub table data here.
                                if bodytext.find(fieldname) >= 0:
                                    templaterows[rownr] = {
                                        "row": row,
                                        "rowtrcopy": copy.deepcopy(row._tr),
                                        "merged": False,
                                        "type": 0,
                                    }
                                    rownr += 1
                                    breakrow = True
                                    break
                                #
                            #
                            if breakrow:
                                break
                        #
                        if breakrow:
                            break
                    #
                    if breakrow:
                        break
                #
                if breakrow:
                    break
            #
        #
        return templaterows

    #

    def tbl_complex_copy(s, p_tblid, p_tbl, p_fields):
        # plpy.notice("TBLID: {}  ,dat\r\n{}".format(p_tblid,p_fields))
        mappedtbl = s.tbl_map.get(p_tblid, None)
        children = None
        limitrow = 0
        distinctdata = {}
        fieldtags = list(p_fields.keys())
        # tblsource = ""
        if mappedtbl is not None:
            children = mappedtbl.get("children")
            limitrow = mappedtbl.get("limitrow", 0)
            # tblsource = mappedtbl.get('nr')
        #

        # sqlmsg("Table Data: {} -> {}  , {}".format(p_tblid,limitrow, mappedtbl  ))
        # if p_fields is None: return
        if p_tbl is None:
            return

        dat = None
        datcnt = 0
        for t in p_fields.keys():
            fielval = p_fields.get(t)
            if fielval is None:
                continue
            dat = fielval.get("value", None)
            if dat is None:
                continue
            if isinstance(dat, dict) or isinstance(dat, list) or isinstance(dat, tuple):
                datcnt = len(dat)
            else:
                continue
            #
            break
        #

        if children is not None:
            for tblid in children:
                childfields = p_fields.get(tblid, {}).get("value", None)
                if childfields is not None and type(childfields) is list:
                    # plpy.notice("Childfields: {}".format(list(childfields[0].keys())))
                    fieldtags.extend(list(childfields[0].keys()))
                #
            #
        #
        if s.tbl_complex_rowdat.get(p_tblid) is None:
            s.tbl_complex_rowdat[p_tblid] = {"min": 0, "max": 0}
        #
        s.tbl_complex_rowdat[p_tblid]["max"] = datcnt
        if datcnt == 0:
            datcnt = 1
        #
        rmrowfuture = {}
        rowdata = s.tbl_complex_headerstart(p_tbl, fieldtags, p_tblid, p_fields)
        # plpy.notice("TBLID: {} len:{},dat\r\n{}\r\nRowdata:{}".format(p_tblid, datcnt,p_fields,rowdata))
        lastrow = None
        for rn in range(
            0, datcnt
        ):  # loop through table data which is the rows for the first column.
            if limitrow > 0 and rn >= limitrow:
                break
            #
            rmrow = False

            # sqlmsg("Table ID: {} Copies: {} rd:{} rmrowfuture:{}".format(p_tblid,rn, len(rowdata), rmrowfuture),"mergetest")
            for (
                rd
            ) in (
                rowdata
            ):  ##this will usually be 1 for 1 row to copy. but can be more than 1 (Generate Blank row feature)
                # plpy.notice("Copies: {} rd:{}".format(rn, rd))

                rmcnt = rmrowfuture.get(str(rn), {}).get("cnt", -1)
                rmrow = False

                if not rowdata[rd]["merged"]:
                    if rowdata[rd].get("type", 0) == 4 and rn != datcnt - 1:
                        continue
                    #
                    lastrow = rowdata[rd]["row"]
                    rowdata[rd]["merged"] = True

                else:
                    if rmcnt >= 0 and rd < rmcnt:
                        rmrow = True
                        pass
                    else:
                        tr = copy.deepcopy(rowdata[rd]["rowtrcopy"])
                        lasttr = lastrow._tr.addnext(tr)
                        lastrow = docx.table._Row(tr, p_tbl)
                    #
                #

                for cell in lastrow.cells:
                    colidx = 0
                    rmcol = False
                    # txt = cell.text
                    # skiptag = txt.find("{}{}:".format(s.TAG_S, s.TAG_ROW_LIMIT))
                    # if skiptag >= 0:
                    #   sqlmsg("Found limit {}".format(txt))
                    #   skiptagend = txt.find("{}".format(s.TAG_E), skiptag)
                    #   if skiptagend +len(s.TAG_S) > skiptag:
                    #     parts = txt[skiptag+len(s.TAG_S):skiptagend].split(":")
                    #     if parts[0] == s.TAG_ROW_LIMIT and parts[1] != "":
                    #       lastrowlimit = int(parts[1])
                    #     #
                    #   #
                    # #

                    for par in cell.paragraphs:
                        for run in par.runs:
                            opdata = s.tag_value_replace(
                                run, par, rn, p_fields, "table", p_tblid, colidx
                            )

                            # remove the blank value row again. Note that its not the filter and depends on value of field with this remove tag.
                            if opdata.get("rmblnk", 0) == 2:
                                rmrow = True
                            if opdata.get("rmblnkcol", 0) == 2:
                                rmcol = True
                            if opdata.get("rmdr", 0) == 2:
                                rmrow = True
                            if opdata.get("limitrow", 0) > 0:
                                limitrow = opdata.get("limitrow", 0)
                            #

                            ##Todo check why we are doing this.
                            # if len(opdata.get("mergedtags", [])) > 0: rmrow = False
                        #
                    #
                    ############ Distinct Rows ##############
                    idxStart = cell.text.find(s.TAG_COND_TBL_DISTINCT)
                    if idxStart >= 0:
                        rowremcnt = 0
                        splitindex = cell.text.find(":", idxStart)
                        splitend = cell.text.find("*]", splitindex)
                        if splitend > 0 and splitend > splitindex:
                            rowremcnt = tryInt(cell.text[splitindex + 1 : splitend])
                        #

                        if distinctdata.get(str(rn), None) is None:
                            distinctdata[str(rn)] = {}
                        #
                        prevvalue = distinctdata.get(str(rn - 1), {}).get(str(colidx))
                        newvalue = cell.text
                        # sqlmsg("Test distinct: prevvalue: {} value: {}".format(prevvalue, newvalue))
                        if newvalue is not None and newvalue != "":
                            distinctdata[str(rn)][str(colidx)] = newvalue
                            if (
                                prevvalue is not None
                                and prevvalue != ""
                                and prevvalue == newvalue
                            ):
                                rmrow = True
                                if rowremcnt > 0:
                                    rmrowfuture[str(rn)] = {
                                        "remove": True,
                                        "cnt": rowremcnt,
                                    }
                                #
                                # rmrowfuture[str(rn+1)] = {"remove": True, "cnt": rowremcnt}
                                # sqlmsg("Found distinct: col:{} row:{} pval:{} nval:{} rowremcnt:{} rmrowfuture:{}".format(colidx,rn,prevvalue,newvalue,rowremcnt,rmrowfuture))
                            #
                        #
                    #
                    ############################################################################################################
                    if (
                        children is not None
                        and cell.tables is not None
                        and len(cell.tables) > 0
                    ):
                        for tblid in children:
                            # plpy.notice("Table in cell: Row {} Tbl nr: {} tblid: {}".format(rn,tblnr, tblid))
                            for tblptr in cell.tables:
                                childfields = p_fields.get(tblid, {}).get("value", None)
                                if (
                                    childfields is not None
                                    and type(childfields) is list
                                    and type(childfields[rn]) is dict
                                ):
                                    # plpy.notice("Found child table: {}  in cell {}".format(tblid, tblptr))
                                    s.tbl_complex_copy(tblid, tblptr, childfields[rn])
                                #

                                # plpy.notice("Child table found: {} {}".format(tblid, tblptr))
                            #
                        #
                    #
                    ############################################################################################################

                    if s.tbl_mode.get(p_tblid) == "free":
                        for tbl in cell.tables:
                            for row2 in tbl.rows:
                                for cell2 in row2.cells:
                                    for par2 in cell2.paragraphs:
                                        for run2 in par2.runs:
                                            s.tag_value_replace(
                                                run2,
                                                par2,
                                                rn,
                                                p_fields,
                                                "table",
                                                p_tblid,
                                                colidx,
                                            )
                                        #
                                    #
                                #
                            #
                        #
                    #

                    if rmcol:
                        s.tbl_rmcol.append(cell)
                    #

                    colidx += 1
                #
                # sqlmsg("row  -> row:{} rm:{}".format(rn,rmrow))
                if rmrow:
                    s.tbl_rmrow.append(lastrow)
                #
            #
        #

    #

    def tbl_complex_init(s):
        if s.inputdata is None:
            return
        if s.tbl_map is None:
            return
        complexfields = s.inputdata.get("complexfields")
        if complexfields is None:
            return

        for nm in s.tbl_map:
            tbldat = s.tbl_map.get(nm)
            if tbldat is None:
                continue
            srctbl = tbldat.get("tblptr")
            # tblid = tbldat.get('nr')
            tblfields = complexfields.get(nm)
            if tblfields is None or srctbl is None:
                plpy.notice(
                    "Table or data [{}] not found for fields -> {}".format(
                        nm, tblfields
                    )
                )
                continue
            #

            s.tbl_complex_copy(nm, srctbl, tblfields)
        #

        for r in s.tbl_rmrow:
            # sqlmsg("Removing row: len: {} -> {}".format(len(s.tbl_rmrow), r) )
            parent = r._tr.find("..")
            if parent is not None:
                parent.remove(r._tr)
                # sqlmsg("Removed row: {}".format(r))
            #
        #

    #

    def globalTagsReplace(s, pRun):
        tag = s.TAG_S + "wordarticle" + s.TAG_E
        tagstart = pRun.text.find(tag)
        if tagstart >= 0:
            line = pRun.text[tagstart:]
            words = line.replace(tag, "").split(" ")
            lookup = 0
            for i in range(0, len(words)):
                if len(words[i]) > 1:
                    lookup = i
                    break
                #
            #
            article = findWordArticle(words[lookup])
            # sqlmsg("Chk wordarticle {} words: {} article: {} line:{}".format(tagstart,words,article,line))
            pRun.text = pRun.text.replace(tag, article)
        #

    #

    def tagErr(s, r, pTagS, pTagE):
        if r.text.find("___debug___") >= 0:
            s.tagdebug = True
            r.text = r.text.replace("___debug___", "")
        #
        i1 = r.text.find(pTagS)
        i2 = r.text.find(pTagE)
        if i2 >= 0 and i1 < 0 or i1 >= 0 and i2 < 0:
            found = False
            tag = ""
            if i1 >= 0:
                tag = r.text[i1:]
            else:
                tag = r.text[:i2]
            #
            for f in s.post_format:
                if f.get("tag", "") == tag and f.get("type", "") == "error":
                    found = True
                    break
                #
            #
            if not found:
                msg = ""
                if s.tagdebug:
                    msg = "Tag must have same styling, copy and paste unformatted."
                #
                s.post_format.append(
                    {"tag": tag, "obj": r, "type": "error", "msg": msg}
                )
            #
        #

    #

    def merge_docx(s):
        basicvalues = {}

        try:
            if len(s.inputdata) <= 0:
                raise MergeToolWarning("No tag data given.")
            #

            s.load_docx()
            doc_styles = s.srcdoc.styles
            try:
                s.white_style = doc_styles["BTHiddenWhite"]
            except Exception as e:
                s.white_style = doc_styles.add_style(
                    "BTHiddenWhite", docx.enum.style.WD_STYLE_TYPE.CHARACTER
                )
                s.white_style.hidden = False
                s.white_style.quick_style = True

                newfont = s.white_style.font
                newfont.color.rgb = RGBColor(255, 255, 255)
            #
            try:
                s.error_style = doc_styles["BTError"]
            except Exception as e:
                s.error_style = doc_styles.add_style(
                    "BTError", docx.enum.style.WD_STYLE_TYPE.CHARACTER
                )
                s.error_style.hidden = False
                s.error_style.quick_style = True

                newfont = s.error_style.font
                newfont.color.rgb = RGBColor(255, 0, 0)
                newfont.size = 12
            #

            def loopbasic_paragraph(p_paragraphs, p_post=False):
                for p in p_paragraphs:
                    if len(p.runs) > 0:
                        for r in p.runs:
                            if p_post:
                                s.globalTagsReplace(r)
                                s.peek_alias_str(r.text, None, "peekdeclare", r, None)
                            else:
                                opdata = s.tag_value_replace(r, p, None)
                                s.tag_picture_add(r)
                                s.peek_alias_str(r.text, None, "peekdeclare", r, None)
                            #
                            s.tagErr(r, s.TAG_S, s.TAG_E)
                            s.tagErr(r, s.TAG_S_EVAL, s.TAG_E_EVAL)
                            s.tagErr(r, s.TAG_S_POSTOP, s.TAG_E_POSTOP)
                        #

                    else:
                        if not p_post:
                            opdata = s.tag_value_replace(p, None, None)
                        #
                    #
                #

            #

            loopbasic_paragraph(s.srcdoc.paragraphs)
            loopbasic_paragraph(s.srcdoc.paragraphs, True)
            # s.replace_shapes(s.srcdoc.inline_shapes)

            if len(s.srcdoc.tables) > 0:
                s.tbl_basic_replace(s.srcdoc.tables)
                s.tbl_complex_init()
                s.tbl_post_loop(s.srcdoc.tables)
            ##

            for p in s.srcdoc.paragraphs:
                if len(p.runs) > 0:
                    for r in p.runs:
                        s.peek_alias_str(r.text, None, "peekdeclare", r, None)
                        s.eval_replace(r, "run")
                    #
                #
            #

            for p in s.srcdoc.paragraphs:
                if len(p.runs) > 0:
                    for r in p.runs:
                        newstr = r.text
                        if (
                            newstr.find(s.TAG_S_ALIASDECL) >= 0
                            or newstr.find(s.TAG_S_ALIASUSE) >= 0
                        ):
                            newstr = s.alias_replace(newstr)
                            r.text = newstr
                        #
                        s.basic_operation_compute(r)
                    #
                #
                if p.text.find(s.TAG_COND_RMP) >= 0:
                    r2.text = r2.text.replace(s.TAG_COND_RMP, "")
                    par.clear()
                #
            #

            def fnTableItemLoop(
                pTableID,
                pTableObj,
                pItemType,
                pItem,
                pIndex=None,
                pRow=None,
                pCell=None,
            ):
                if pItemType == "run":
                    newstr = s.alias_replace(pItem.text, None, pIndex)
                    if newstr != pItem.text:
                        pItem.text = newstr
                    #

                    s.basic_operation_compute(pItem, pTableID)
                    s.tagErr(pItem, s.TAG_S, s.TAG_E)
                    s.tagErr(pItem, s.TAG_S_EVAL, s.TAG_E_EVAL)
                    s.tagErr(pItem, s.TAG_S_POSTOP, s.TAG_E_POSTOP)
                #

            #

            s.visit_table_cb(s.srcdoc.tables, 0, fnTableItemLoop)

            s.alias_clean()

            ## Loop to appy formatting
            for i in s.post_format:
                tag = i.get("tag", "")
                obj = i.get("obj", None)
                ftype = i.get("type", "esign")
                msg = i.get("msg", "")
                if isinstance(obj, docx.text.run.Run):
                    # sqlmsg("Debug post_format --> type:{} tag:{} objtype:{} text:{} style:{}".format(ftype, tag, type(obj), obj.text, s.white_style))
                    # newfont = copy.deepcopy(obj.font)
                    # newfont.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    # obj.font = newfont
                    if ftype == "error":
                        obj.style = s.error_style
                    elif ftype == "remove":
                        obj.style = None
                    else:
                        obj.style = s.white_style
                        # obj.text = obj.text + ' styled'
                    #
                    if msg != "":
                        obj.text = obj.text + "[" + msg + "]"
                    #
                #
            #
            s.docx_textboxes_lp("merge")
            ###Header/Footer Basic Text
            # s.header_replace(s.inputdata.get('fields'))
            # s.footer_replace(s.inputdata.get('fields'))

            # basics
            s.actionHeadFootTags("head", "merge", s.inputdata.get("fields"))
            s.actionHeadFootTags("foot", "merge", s.inputdata.get("fields"))
            # tables
            s.actionHeadFootTags("head", "merge", s.inputdata.get("fields"), True)
            s.actionHeadFootTags("foot", "merge", s.inputdata.get("fields"), True)
            for objlist in s.cleanupList:
                obj = objlist.get("obj", None)
                tag = objlist.get("tag", "")

                if isinstance(obj, docx.text.run.Run):
                    if obj.text is not None and obj.text.find(tag) >= 0:
                        if (
                            s.cleanMissing
                            and not tag.lower().find(s.STR_TAG_ESIGN) >= 0
                        ):
                            obj.text = obj.text.replace(tag, "")
                            plpy.notice("Orphan tag to be removed, {}".format(objlist))
                        #
                    #
                    # plpy.notice("Orphan removed, {}".format(obj.text))
                #
            #

            # sqlmsg(json.dumps(s.aliasList), "aliasList")

            s.save_docx()

        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            raise MergeToolError("Failed merging docx: " + str(e), "merge_docx", exc_tb)
        #

    #

    def values_exists(s, p_obj, p_fields):
        for tag in s.inputdata:
            if type(s.inputdata.get(tag)) is not dict:
                raise MergeToolError(
                    "Expecting tag field to have a type and value property.",
                    "tag_value_replace",
                )
            #
            tagtype = int(s.inputdata.get(tag).get("type", 0))
            tagvalue = s.inputdata.get(tag).get("value", 0)

            if p_str.find(tag) >= 0:
                if type(tagvalue) is list:
                    return {"table": "", "field": tag, "values": tagvalue}
                #
            #
        #
        return None

    #

    def fixanchor_to_inline(s, p_filename):
        """This function read a docx file and replaces anchor tag images."""
        import zipfile
        import lxml.etree as ET

        global docstr
        docstr = None

        def updateZip(p_filename, p_filename2, p_replacefilename, p_data):
            # create a temp copy of the archive without filename
            with zipfile.ZipFile(p_filename, "r") as zin:
                with zipfile.ZipFile(p_filename2, "w") as zout:
                    zout.comment = zin.comment  # preserve the comment
                    for item in zin.infolist():
                        if item.filename == p_replacefilename:
                            zout.writestr(item, p_data)
                            plpy.notice(
                                "Replacing file in zip {}, {}".format(
                                    item, p_replacefilename
                                )
                            )
                        else:
                            zout.writestr(item, zin.read(item.filename))
                        #

            # replace with the temp archive
            os.remove(p_filename)
            os.rename(p_filename2, p_filename)
            #

        #

        def replaceexml():
            global docstr
            NSMAP = {
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
                "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
                "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
            }

            tree = ET.fromstring(docstr)
            for e in tree.findall("w:body/w:p/w:r/w:drawing/wp:anchor", NSMAP):
                docpt = e.find("wp:docPr", NSMAP)
                graphic = e.find("a:graphic", NSMAP)
                ##Yeah, we need to check if there is a graphic within and filename, else we brake stuff when replacing.
                if docpt is not None and graphic is not None:
                    name = docpt.get("name")
                    filename = docpt.get("descr")
                    if filename != "":
                        e.tag = ET.QName(
                            "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
                            "inline",
                        )
                        plpy.notice(
                            "Found replaceable file: {}, {}".format(name, filename)
                        )
                    #
                #
                # plpy.notice("Found anchor:")
            #
            docstr = ET.tostring(tree)

        #

        if not zipfile.is_zipfile(p_filename):
            raise Exception("Not a valid docx zip file.")
        #

        with zipfile.ZipFile(p_filename, "a") as zp:
            docinfo = zp.getinfo("word/document.xml")
            with zp.open(docinfo, "r") as docx:
                docstr = docx.read()
                plpy.notice("Read: {}".format(len(docstr)))
            #
        #

        replaceexml()
        # plpy.notice(docstr.decode('utf8'))
        updateZip(p_filename, p_filename + ".temp.zip", "word/document.xml", docstr)

    #

    def generate_mockup(s):
        s.load_docx()

        lastsection = ""
        p = s.srcdoc.add_paragraph("")
        p.style = s.srcdoc.styles["Heading 1"]
        r = p.add_run("Mergefield Test Document (mocktest_docx):")
        r = p.add_run(
            "\r\nGenerated: {0:%Y-%m-%d %H:%M:%S}".format(datetime.datetime.now())
        )
        lasttable = ""
        objtable = None

        for tag in s.inputdata:
            if type(s.inputdata.get(tag)) is not dict:
                raise MergeToolError(
                    "Expecting tag field to have a type and value property.",
                    "tag_value_replace",
                )
            #
            tagtype = int(s.inputdata.get(tag).get("type", 0))
            tooltip = s.inputdata.get(tag).get("tooltip", "")
            display_name = s.inputdata.get(tag).get("display_name", "")
            parent_name = s.inputdata.get(tag).get("parent_name", "")
            parent_tooltip = s.inputdata.get(tag).get("parent_tooltip", "")
            parent_table_name = s.inputdata.get(tag).get("parent_table", "")
            table_name = s.inputdata.get(tag).get("table_name", "")
            if table_name == "":
                table_name = parent_table_name

            if parent_name != lastsection:
                lastsection = parent_name
                p = s.srcdoc.add_paragraph("\r\n")
                r = p.add_run("{}\r\n".format(parent_name))
                r.bold = True
                r = p.add_run("Tooltip:{}\r\n".format(parent_tooltip))
                r.bold = False
                r.italic = True
            #
            if tagtype in (0, 1, 3):
                p = s.srcdoc.add_paragraph("")
                r = p.add_run(
                    "{}:\r\nTag:{}\r\n".format(
                        display_name, tag.strip(s.TAG_S).strip(s.TAG_E)
                    )
                )
                r.bold = True
                r = p.add_run("Tooltip:{}\r\n".format(tooltip))
                r.bold = False
                r.italic = True
                r = p.add_run("Data: [{}]\r\n".format(tag))
                r.bold = False
            elif tagtype == 2:
                if lasttable != parent_name:
                    lasttable = parent_name
                    plpy.notice(
                        "new table: {}, {}, {}".format(parent_name, display_name, tag)
                    )
                    # if objtable
                #

                plpy.notice(
                    "last table: {}, {}, {}, {}".format(
                        table_name, parent_table_name, display_name, tag
                    )
                )
            #

        #
        s.save_docx()

    #

    def splitRun(s, p_run, p_par, p_startstr, p_endstr):
        ##todo complete
        runtext = p_run.text
        start = runtext.find(p_startstr)
        end = runtext.find(p_endstr)
        if end > start and start > 0:
            for r in p_par._p.r_lst:
                if r == p_run._r:
                    # sqlmsg("Found run: {}", r)
                    p_run._r = Run(r, p_par)
                    break
                #
            #
        #

    #

    def placeDoc(s, p_obj, p_parent, tag, val):
        ### Places document inside exiting object.
        if not os.path.exists(val):
            raise Exception(
                "Merge tag file does not exist {} Mergetag: {}".format(val, tag)
            )
        if p_obj is None:
            raise Exception("Null object given. Mergtag: {}".format(tag))

        par = p_parent
        if str(type(p_obj)).find("Paragraph") >= 0:
            par = p_obj
            # sqlmsg("placeDoc paragraph only")
        #

        tagdoc = Document(val)

        def placeTable(p_table, p_destpar):
            tbl, p = p_table._tbl, p_destpar._p
            new_tbl = copy.deepcopy(tbl)
            p.addnext(new_tbl)

        #

        for p in tagdoc.paragraphs:
            # sqlmsg("placeDoc placing {} into {}".format(p.text, par.text))
            new_p = copy.deepcopy(p._p)
            new_p.style = None
            par._p.addnext(new_p)
            # newpar = par.insert_paragraph_before("Test",p.style)
            # for r in p.runs:
            #  nr = newpar.add_run(r.text)
            #  nr.bold = r.bold
            #  nr.italic = r.italic
            #  nr.underline = r.underline
            #
        #

        if len(tagdoc.tables) > 0:
            for tbl in tagdoc.tables:
                placeTable(tbl, par)
            #
        #

    #


#
